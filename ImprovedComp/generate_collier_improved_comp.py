"""
Improved Comparable Generator — Collier PDF Source
Parses the Collier appraisal-format improved sale comps PDF and generates:
  - ImprovedComp_Collier_CStore.xlsx   (38-field comparison spreadsheet)
  - ImprovedComp_Collier_CStore.docx   (Word doc with data table + PDF page images)

Usage:
    cd c:\\AppraisalWorkspace\\ImprovedComp
    python generate_collier_improved_comp.py

    Or with CLI args:
    python generate_collier_improved_comp.py --name "CStore" --pdf "path/to/collier.pdf"

Source PDF format: Colliers International Valuation & Advisory Services
  - Each page contains one comparable sale
  - Structured fields: Location Info, Sale Info, Physical Info, Operating Income,
    Analysis Info, Confirmation, Remarks
"""

import os
import re
import io
import argparse
import tempfile
import fitz  # PyMuPDF
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ============================================================
# CONFIG
# ============================================================
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
SOURCES_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), "Sources")

PDF_PATH    = os.path.join(SOURCES_DIR, "Comps", "improved cstore comps type up collier.pdf")
OUTPUT_DIR  = os.path.join(SCRIPT_DIR, "Output")
OUTPUT_NAME = "Collier_CStore"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Image rendering
RENDER_DPI      = 200
IMAGE_WIDTH_IN  = 6.5
JPEG_QUALITY    = 85

# ============================================================
# HELPERS
# ============================================================

def _clean(s):
    if not s:
        return ""
    return s.replace("\xa0", " ").replace("\u00b0", "").strip()

def _money_fmt(raw_str):
    s = str(raw_str).replace(",", "").replace("$", "").strip()
    if not s:
        return ""
    try:
        v = float(s)
        return f"${v:,.0f}" if v == int(v) else f"${v:,.2f}"
    except ValueError:
        return ""

# ============================================================
# COLLIER PDF PARSER
# ============================================================

def _extract_field(text, label):
    """Extract value that follows a label on the next line (Collier format)."""
    pattern = re.escape(label) + r"\s*\n\s*(.+)"
    m = re.search(pattern, text)
    return _clean(m.group(1)) if m else ""

def _extract_inline(text, label):
    """Extract value on same line as label."""
    pattern = re.escape(label) + r"\s+(.+)"
    m = re.search(pattern, text)
    return _clean(m.group(1)) if m else ""

def parse_collier_page(page_text):
    """Parse one Collier comparable page into a raw dict."""
    raw = {}

    # Comparable number
    comp_m = re.search(r"COMPARABLE\s+(\d+)", page_text)
    raw["comp_number"] = int(comp_m.group(1)) if comp_m else 0

    # Location Information
    raw["name"] = _extract_field(page_text, "Name")
    raw["address"] = _extract_field(page_text, "Address")

    csz = _extract_field(page_text, "City, State, Zip Code")
    if csz:
        csz_m = re.match(r"(.+?),\s*([A-Z]{2}),?\s*(\d{5})", csz)
        if csz_m:
            raw["city"] = csz_m.group(1).strip()
            raw["state"] = "Texas"
            raw["zip"] = csz_m.group(3)
        else:
            raw["city"] = csz
            raw["state"] = "Texas"

    raw["county"] = _extract_field(page_text, "County")
    raw["msa"] = _extract_field(page_text, "MSA")
    raw["apn"] = _extract_field(page_text, "APN")

    # Sale Information
    raw["buyer"] = _extract_field(page_text, "Buyer")
    raw["seller"] = _extract_field(page_text, "Seller")
    raw["transaction_date"] = _extract_field(page_text, "Transaction Date")
    raw["transaction_status"] = _extract_field(page_text, "Transaction Status")

    tp = _extract_field(page_text, "Transaction Price")
    raw["transaction_price"] = tp.replace("$", "").replace(",", "") if tp else ""

    ap = _extract_field(page_text, "Analysis Price")
    raw["analysis_price"] = ap.replace("$", "").replace(",", "") if ap else ""

    raw["recording_number"] = _extract_field(page_text, "Recording Number")
    raw["rights_transferred"] = _extract_field(page_text, "Rights Transferred")
    raw["financing"] = _extract_field(page_text, "Financing")
    raw["conditions_of_sale"] = _extract_field(page_text, "Conditions of Sale")

    # Physical Information
    gba = _extract_field(page_text, "Gross Building Area (GBA)")
    raw["gba"] = gba.replace(",", "") if gba else ""

    nra = _extract_field(page_text, "Leasable Area (NRA)")
    raw["nra"] = nra.replace(",", "") if nra else ""

    raw["year_built"] = _extract_field(page_text, "Year Built")
    raw["year_renovated"] = _extract_field(page_text, "Year Renovated")
    raw["front_footage"] = _extract_field(page_text, "Front Footage")
    raw["quality"] = _extract_field(page_text, "Quality")
    raw["condition"] = _extract_field(page_text, "Condition")
    raw["appeal"] = _extract_field(page_text, "Appeal")
    raw["building_structure"] = _extract_field(page_text, "Building Structure")
    raw["exterior"] = _extract_field(page_text, "Exterior")

    # Site Size — "0.5 Acres (20,996 SF)"
    site = _extract_field(page_text, "Site Size")
    if site:
        site_m = re.match(r"([\d.]+)\s*Acres?\s*\(([\d,]+)\s*SF\)", site)
        if site_m:
            raw["site_acres"] = site_m.group(1)
            raw["site_sf"] = site_m.group(2).replace(",", "")
        else:
            raw["site_acres"] = ""
            raw["site_sf"] = ""
    raw["zoning"] = _extract_field(page_text, "Zoning")
    raw["topography"] = _extract_field(page_text, "Topography")
    raw["access"] = _extract_field(page_text, "Access")

    # Operating Income
    noi_m = re.search(r"Net Operating Income\s+\$([\d,]+)\s+\$([\d.]+)", page_text)
    if noi_m:
        raw["noi_total"] = noi_m.group(1).replace(",", "")
        raw["noi_per_sf"] = noi_m.group(2)

    occ = _extract_field(page_text, "Occupancy at Sale")
    raw["occupancy"] = occ

    # Analysis Information
    ppsf = _extract_field(page_text, "Price per SF")
    raw["price_per_sf"] = ppsf.replace("$", "").replace(",", "") if ppsf else ""

    adj_ppsf = _extract_field(page_text, "Adjusted Price per SF")
    raw["adjusted_price_per_sf"] = adj_ppsf.replace("$", "").replace(",", "") if adj_ppsf else ""

    cap_m = re.search(r"Capitalization Rate\s+([\d.]+)%", page_text)
    raw["cap_rate"] = cap_m.group(1) + "%" if cap_m else ""

    site_cov = _extract_field(page_text, "Site Coverage (SF)/Ratio")
    raw["site_coverage"] = site_cov

    # Confirmation
    raw["confirmation_source"] = _extract_field(page_text, "Source")

    # Exposure
    exposure = _extract_field(page_text, "Exposure")
    raw["exposure"] = exposure

    # Remarks
    remarks_m = re.search(r"REMARKS\s*\n(.*?)$", page_text, re.DOTALL)
    raw["remarks"] = _clean(remarks_m.group(1).replace("\n", " ")) if remarks_m else ""

    return raw


def parse_collier_pdf(pdf_path):
    """Parse a Collier improved comps PDF. Returns list of (raw_dict, page_index) tuples."""
    pdf = fitz.open(pdf_path)
    results = []
    for i in range(pdf.page_count):
        text = pdf[i].get_text()
        if "COMPARABLE" in text:
            raw = parse_collier_page(text)
            raw["_page_index"] = i
            results.append(raw)
    pdf.close()
    return results

# ============================================================
# FIELD MAPPING — same 38-field structure as ImprovedComp
# ============================================================

FIELD_ORDER = [
    "Comparable Sale",
    "Property Type",
    "CAD ID:",
    "Street Address:",
    "City:",
    "State:",
    "County:",
    "Land Size (SF):",
    "Land Size (Acres):",
    "Gross Building Area (SF):",
    "Net Rentable Area (SF)",
    "Rentable Unit Number",
    "Average Unit Size",
    "Unit Mix",
    "Project Amenities",
    "Unit Amenities",
    "Finish-out Percentage",
    "Land to Building Ratio:",
    "Year Built/ Renovated",
    "Date of Sale:",
    "Sales Price:",
    "Unit Price ($/Net SF):",
    "Unit Price ($/Gross SF):",
    "Unit Price ($/Unit)",
    "Cap Rate:",
    "NOI ($/Net SF)",
    "NOI ($/Unit)",
    "Occupancy Rate at Time of Sale:",
    "Data Source:",
    "Recorded  Number:",
    "Grantor:",
    "Grantee:",
    "Terms and Conditions:",
    "Time on Market:",
    "Property Rights Conveyed:",
    "Transactional Status",
    "Original Listing Price:",
    "Sale-to-List Ratio:",
    "Additional Comments:",
]

def derive_fields(raw, sale_number=1):
    """Map Collier raw fields to the 38-field ImprovedComp structure."""
    d = {}

    d["Comparable Sale"] = f"Sale No. {sale_number}"
    d["Property Type"] = "Convenience Store / Gas Station"
    d["CAD ID:"] = raw.get("apn", "")
    d["Street Address:"] = raw.get("address", "")
    d["City:"] = raw.get("city", "")
    d["State:"] = raw.get("state", "Texas")
    d["County:"] = raw.get("county", "")

    site_sf = raw.get("site_sf", "")
    site_acres = raw.get("site_acres", "")
    gba = raw.get("gba", "")
    nra = raw.get("nra", "")

    d["Land Size (SF):"] = site_sf
    d["Land Size (Acres):"] = site_acres
    d["Gross Building Area (SF):"] = gba
    d["Net Rentable Area (SF)"] = nra

    # These fields are not in the Collier format — leave blank for analyst
    d["Rentable Unit Number"] = ""
    d["Average Unit Size"] = ""
    d["Unit Mix"] = ""
    d["Project Amenities"] = ""
    d["Unit Amenities"] = ""
    d["Finish-out Percentage"] = ""

    # Land to Building Ratio
    if site_sf and gba:
        try:
            d["Land to Building Ratio:"] = f"{float(site_sf) / float(gba):.2f}"
        except (ValueError, ZeroDivisionError):
            d["Land to Building Ratio:"] = ""
    else:
        d["Land to Building Ratio:"] = ""

    # Year Built / Renovated
    yr = raw.get("year_built", "")
    yr_ren = raw.get("year_renovated", "")
    if yr and yr_ren:
        d["Year Built/ Renovated"] = f"{yr} / {yr_ren}"
    else:
        d["Year Built/ Renovated"] = yr

    d["Date of Sale:"] = raw.get("transaction_date", "")

    # Use analysis price (adjusted) as the sale price if available
    price_raw = raw.get("analysis_price", "") or raw.get("transaction_price", "")
    d["Sales Price:"] = _money_fmt(price_raw)

    # Price per SF
    if price_raw and nra:
        try:
            d["Unit Price ($/Net SF):"] = f"${float(price_raw) / float(nra):,.2f}"
        except (ValueError, ZeroDivisionError):
            d["Unit Price ($/Net SF):"] = ""
    else:
        d["Unit Price ($/Net SF):"] = raw.get("price_per_sf", "")

    if price_raw and gba:
        try:
            d["Unit Price ($/Gross SF):"] = f"${float(price_raw) / float(gba):,.2f}"
        except (ValueError, ZeroDivisionError):
            d["Unit Price ($/Gross SF):"] = ""
    else:
        d["Unit Price ($/Gross SF):"] = ""

    d["Unit Price ($/Unit)"] = ""

    d["Cap Rate:"] = raw.get("cap_rate", "")

    noi_per_sf = raw.get("noi_per_sf", "")
    d["NOI ($/Net SF)"] = f"${noi_per_sf}" if noi_per_sf else ""
    d["NOI ($/Unit)"] = ""

    d["Occupancy Rate at Time of Sale:"] = raw.get("occupancy", "")

    source = raw.get("confirmation_source", "")
    d["Data Source:"] = f"Colliers International - {source}" if source else "Colliers International"

    d["Recorded  Number:"] = raw.get("recording_number", "")
    d["Grantor:"] = raw.get("seller", "")
    d["Grantee:"] = raw.get("buyer", "")

    financing = raw.get("financing", "")
    if "Cash" in financing:
        d["Terms and Conditions:"] = "Cash at Settlement"
    else:
        d["Terms and Conditions:"] = financing

    d["Time on Market:"] = ""  # Not available in Collier format
    d["Property Rights Conveyed:"] = raw.get("rights_transferred", "")
    d["Transactional Status"] = raw.get("transaction_status", "")
    d["Original Listing Price:"] = ""  # Not available in Collier format
    d["Sale-to-List Ratio:"] = ""  # Not available

    d["Additional Comments:"] = raw.get("remarks", "")

    return d

# ============================================================
# EXCEL GENERATION (same format as existing ImprovedComp)
# ============================================================

def _border(style="thin"):
    s = Side(border_style=style, color="000000")
    return Border(left=s, right=s, top=s, bottom=s)

def generate_excel(all_comps, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    ws.column_dimensions["A"].width = 2
    ws.column_dimensions["B"].width = 32
    for col in ["C", "D", "E", "F", "G"]:
        ws.column_dimensions[col].width = 24

    hdr_fill   = PatternFill("solid", fgColor="4472C4")
    label_fill = PatternFill("solid", fgColor="D9E1F2")
    alt_fill   = PatternFill("solid", fgColor="F2F2F2")
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    hdr_font   = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    label_font = Font(name="Calibri", size=10)
    val_font   = Font(name="Calibri", size=10)
    center_al  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_al    = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    # Row 1 — header
    ws.cell(row=1, column=2).value     = "Field"
    ws.cell(row=1, column=2).font      = hdr_font
    ws.cell(row=1, column=2).fill      = hdr_fill
    ws.cell(row=1, column=2).alignment = center_al
    ws.cell(row=1, column=2).border    = _border()

    for col_idx, comp in enumerate(all_comps, start=3):
        cell = ws.cell(row=1, column=col_idx)
        cell.value     = comp.get("Comparable Sale", f"Sale No. {col_idx - 2}")
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = center_al
        cell.border    = _border()

    # Data rows
    for row_idx, field in enumerate(FIELD_ORDER, start=2):
        use_alt  = (row_idx % 2 == 0)
        row_fill = alt_fill if use_alt else white_fill

        label_cell = ws.cell(row=row_idx, column=2)
        label_cell.value     = field.rstrip(":").rstrip()
        label_cell.font      = label_font
        label_cell.fill      = label_fill
        label_cell.alignment = left_al
        label_cell.border    = _border()

        for ci, comp in enumerate(all_comps):
            val_cell = ws.cell(row=row_idx, column=ci + 3)
            val_cell.font      = val_font
            val_cell.fill      = row_fill
            val_cell.alignment = left_al
            val_cell.border    = _border()

            val = comp.get(field, "")
            # Write numeric where appropriate
            for numeric_field in ("Land Size (SF):", "Gross Building Area (SF):",
                                  "Net Rentable Area (SF)"):
                if field == numeric_field and val:
                    try:
                        val_cell.value = int(val)
                        break
                    except (ValueError, TypeError):
                        pass
            else:
                val_cell.value = val

    ws.freeze_panes = "B2"
    wb.save(output_path)
    print(f"  Excel saved: {output_path}")

# ============================================================
# PAGE IMAGE RENDERING (JPEG, no orphaned blobs)
# ============================================================

def render_page_as_jpeg(pdf_path, page_index):
    """Render a single PDF page as compressed JPEG bytes."""
    pdf = fitz.open(pdf_path)
    page = pdf[page_index]
    scale = RENDER_DPI / 72.0
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)

    # Convert to JPEG via PIL
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=JPEG_QUALITY)
    buf.seek(0)
    jpeg_bytes = buf.read()

    pdf.close()
    return jpeg_bytes, pix.width, pix.height

# ============================================================
# WORD DOC GENERATION
# ============================================================

def _add_data_table(doc, comp):
    """Add data table showing only non-blank fields."""
    visible = [(f, comp.get(f, "")) for f in FIELD_ORDER if comp.get(f, "")]
    if not visible:
        return

    table = doc.add_table(rows=len(visible), cols=2)
    table.style = "Table Grid"

    for row_idx, (field, value) in enumerate(visible):
        row = table.rows[row_idx]

        label_cell = row.cells[0]
        label_cell.text = field.rstrip(":").rstrip()
        lp = label_cell.paragraphs[0]
        if lp.runs:
            lp.runs[0].font.bold = True
            lp.runs[0].font.size = Pt(9)
            lp.runs[0].font.name = "Calibri"
        label_cell.width = Inches(2.5)

        val_cell = row.cells[1]
        val_cell.text = str(value)
        vp = val_cell.paragraphs[0]
        if vp.runs:
            vp.runs[0].font.size = Pt(9)
            vp.runs[0].font.name = "Calibri"
        val_cell.width = Inches(3.8)

    return table

def generate_docx(all_comps, pdf_path, output_path):
    """Generate Word doc with data tables + PDF page images (JPEG)."""
    doc = Document()

    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    for sale_idx, comp in enumerate(all_comps, start=1):
        # Heading
        heading = doc.add_paragraph(style="Normal")
        run = heading.add_run(f"Improved Comparable Sale No. {sale_idx}")
        run.bold      = True
        run.font.size = Pt(12)
        run.font.name = "Calibri"

        # Data table
        _add_data_table(doc, comp)
        doc.add_paragraph()  # spacer

        # PDF page image (rendered as JPEG)
        page_idx = comp.get("_page_index")
        if page_idx is not None:
            jpeg_bytes, w, h = render_page_as_jpeg(pdf_path, page_idx)
            target_w = Inches(IMAGE_WIDTH_IN)
            target_h = int(target_w * (h / w))
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(io.BytesIO(jpeg_bytes), width=target_w, height=Emu(target_h))
            doc.add_paragraph()  # spacer

    doc.save(output_path)
    print(f"  Word doc saved: {output_path}")

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate ImprovedComp Excel + Word doc from Collier appraisal PDF"
    )
    parser.add_argument("--name", help="Output name (e.g. 'CStore')")
    parser.add_argument("--pdf",  help="Collier improved comps PDF path")
    args = parser.parse_args()

    if args.name:
        OUTPUT_NAME = args.name
    if args.pdf:
        PDF_PATH = args.pdf

    if not os.path.exists(PDF_PATH):
        print(f"ERROR: PDF not found: {PDF_PATH}")
        raise SystemExit(1)

    print(f"Parsing Collier PDF: {PDF_PATH}")
    comps_raw = parse_collier_pdf(PDF_PATH)
    print(f"  Found {len(comps_raw)} comparable(s)")

    all_comps = []
    for i, raw in enumerate(comps_raw, start=1):
        comp = derive_fields(raw, sale_number=i)
        comp["_page_index"] = raw["_page_index"]
        all_comps.append(comp)
        print(f"  Sale {i}: {raw.get('name','?')} - {raw.get('address','?')}, {raw.get('city','?')}")
        print(f"    Price: {comp.get('Sales Price:','')}  GBA: {raw.get('gba','')}  "
              f"Cap: {raw.get('cap_rate','')}  NOI/SF: {raw.get('noi_per_sf','')}")

    print("\nGenerating Excel...")
    xlsx_path = os.path.join(OUTPUT_DIR, f"ImprovedComp_{OUTPUT_NAME}.xlsx")
    generate_excel(all_comps, xlsx_path)

    print("\nGenerating Word doc...")
    docx_path = os.path.join(OUTPUT_DIR, f"ImprovedComp_{OUTPUT_NAME}.docx")
    generate_docx(all_comps, PDF_PATH, docx_path)

    print(f"\nDone. Output in: {OUTPUT_DIR}")
