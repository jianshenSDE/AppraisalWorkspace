"""
Improved Comparable Generator — generate_improved_comp.py

Parses one or more MLS PDFs (+ optional CAD PDFs) to generate:
  - ImprovedComp_[name].xlsx   (38-field comparison spreadsheet)
  - ImprovedComp_[name].docx   (Word doc with data table + MLS photos)

Usage — edit CONFIG below, then run:
    cd c:\\AppraisalWorkspace\\ImprovedComp
    python generate_improved_comp.py

Usage — command-line arguments:
    python generate_improved_comp.py --name "Terrell" --mls "path/MLS.pdf" --cad "path/CAD.pdf"

    For multiple comparables:
    python generate_improved_comp.py --name "DFW" --mls "MLS1.pdf" --mls "MLS2.pdf" --cad "CAD1.pdf" --cad "CAD2.pdf"

    Sources:
      --mls        MLS PDF export (required)
      --cad        County Appraisal District PDF (recommended — provides bldg SF, year built, county)
      --texasfile  TexasFile.com screenshot PNG (optional alternative)
"""

import os
import re
import argparse
import tempfile
import fitz  # PyMuPDF
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONFIG — Update these for each new report
# ============================================================
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
SOURCES_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), "Sources")

MLS_PDF_PATHS = [
    # os.path.join(SOURCES_DIR, "MLS", "808 E Moore Avenue, Terrell MLS.pdf"),
]

CAD_PDF_PATHS = [
    # os.path.join(SOURCES_DIR, "CAD", "808 E Moore Avenue, Terrell Kaufmann CAD.pdf"),
]

TEXASFILE_PNG_PATHS = [
    # os.path.join(SOURCES_DIR, "TexasFile", "texasfile-search.png"),
]

OUTPUT_DIR  = os.path.join(SCRIPT_DIR, "Output")
OUTPUT_NAME = "Output"  # used in output filenames

# ============================================================
# HELPERS
# ============================================================

def _clean(s):
    if not s:
        return ""
    return s.replace("\xa0", " ").replace("\u00b0", "").strip()

def _first(pattern, text, group=1, flags=re.IGNORECASE | re.MULTILINE):
    m = re.search(pattern, text, flags)
    return _clean(m.group(group)) if m else ""

def _money_fmt(raw_str):
    s = str(raw_str).replace(",", "").replace("$", "").strip()
    if not s:
        return ""
    try:
        v = float(s)
        return f"${v:,.0f}" if v == int(v) else f"${v:,.2f}"
    except ValueError:
        return ""

def _pct_fmt(val):
    """0.88 -> '88%', 1 -> '100%'"""
    try:
        return f"{float(val):.0%}"
    except (ValueError, TypeError):
        return str(val) if val else ""

def _flip_name(name):
    """'Perez Fernando' -> 'Fernando Perez', 'Wallen Joshua & Brandi' -> 'Joshua and Brandi Wallen'"""
    name = _clean(name)
    if not name:
        return ""
    if "&" in name:
        parts = [p.strip() for p in name.split("&")]
        first_parts = parts[0].split()
        if len(first_parts) >= 2:
            last   = first_parts[0]
            first1 = " ".join(first_parts[1:])
            first2 = parts[1].strip()
            return f"{first1} and {first2} {last}"
        return name
    parts = name.split()
    if len(parts) >= 2:
        return " ".join(parts[1:]) + " " + parts[0]
    return name

def _format_terms(financing):
    f = financing.lower().strip()
    if "cash" in f:
        return "Cash or conventional financing"
    if "conventional" in f:
        return "Conventional financing"
    if "fha" in f:
        return "FHA financing"
    if "va" in f:
        return "VA financing"
    return financing.strip().title() if financing.strip() else ""

# ============================================================
# PROPERTY TYPE DETECTION
# ============================================================

PROPERTY_TYPE_KEYWORDS = [
    (["taqueria", "restaurant", "bar ", "grill", "tavern", "diner", "food", "eatery"], "Restaurant"),
    (["retail", "store", "shop", "boutique", "strip center", "strip mall"],            "Retail"),
    (["office", "professional", "suite", "corporate"],                                  "Office"),
    (["warehouse", "storage", "industrial", "flex", "distribution", "manufacturing"],   "Industrial/Warehouse"),
    (["medical", "clinic", "dental", "health", "pharmacy"],                             "Medical Office"),
    (["hotel", "motel", "hospitality", "inn", "lodging"],                               "Hospitality"),
    (["church", "worship", "ministry", "chapel", "faith"],                              "Religious"),
    (["auto", "car wash", "tire", "mechanic", "dealership"],                            "Automotive"),
    (["daycare", "childcare", "preschool"],                                              "Childcare"),
    (["bank", "financial", "credit union"],                                              "Bank/Financial"),
]

def _detect_property_type(description_text):
    low = description_text.lower()
    for keywords, prop_type in PROPERTY_TYPE_KEYWORDS:
        if any(kw in low for kw in keywords):
            return prop_type
    return ""

def _detect_occupancy(description_text):
    low = description_text.lower()
    if any(w in low for w in ["currently leased", "tenant in place", "leased out",
                               "occupied", "established tenant", "existing tenant"]):
        return "100%"
    if any(w in low for w in ["vacant", "empty", "unoccupied", "available for lease"]):
        return "0%"
    return ""

# ============================================================
# MLS PDF PARSER
# ============================================================

def _parse_listing_text(full_text, p1):
    raw = {}

    # Address
    addr_m = re.search(
        r"[^\w\n]*([^\n,]+),\s*([^,\n]+),\s*Texas\s+(\d{5})",
        p1, re.IGNORECASE
    )
    if addr_m:
        raw["street"] = _clean(addr_m.group(1))
        raw["city"]   = _clean(addr_m.group(2))
        raw["state"]  = "Texas"

    raw["mls_number"]  = _first(r"MLS#:\s*(\d+)", full_text)
    raw["county"]      = _first(r"County:\s*\n?\s*([A-Za-z][A-Za-z ]*?)(?:\n|Lake Name|$)", full_text)
    raw["cad_id"]      = _first(r"Parcel ID:\s*\n?\s*(\d+)", full_text)
    raw["lot_sf"]      = _first(r"Lot SqFt:\s*([\d,]+)", full_text).replace(",", "")
    raw["acres"]       = _first(r"(?:^|[\s\n])Acres:\s*\n?\s*([\d.]+)", full_text)
    raw["bldg_sf"]     = _first(r"Building Sq Ft:\s*([\d,]+)", full_text).replace(",", "")
    raw["close_price"] = _first(r"Close Price:\s*\$([\d,]+)", full_text).replace(",", "")
    raw["close_date"]  = _first(r"Closed Date:\s*(\d{1,2}/\d{1,2}/\d{4})", full_text)

    lp = _first(r"OLP:\s*\$([\d,]+)", full_text).replace(",", "")
    if not lp:
        lp = _first(r"(?:^|\n)LP:\s*\$?\s*\n?\s*\$?([\d,]+)", full_text).replace(",", "")
    raw["list_price"] = lp

    cdom = _first(r"\bCDOM:\s*(\d+)", full_text)
    dom  = _first(r"\bDOM:\s*(\d+)", full_text)
    raw["dom"]      = cdom or dom
    raw["financing"] = _first(r"Buyer Financing:\s*([^\n]+)", full_text)

    if "NTREIS" in full_text:
        raw["mls_name"] = "North Texas MLS"
    elif "HAR" in full_text:
        raw["mls_name"] = "Houston MLS"
    else:
        raw["mls_name"] = "MLS"

    # Property description (for type detection and comments)
    desc_m = re.search(
        r"(?:Property Description|Public Remarks|Remarks|Comments):\s*\n?(.*?)(?:\nPublic Driving|\nAgent/Office|\nFinancial|\nShowing|\Z)",
        full_text, re.DOTALL | re.IGNORECASE
    )
    raw["description"] = _clean(desc_m.group(1).replace("\n", " ")) if desc_m else ""

    # Fallback: run keyword detection on full text so type/occupancy are never missed
    raw["property_type_hint"] = _detect_property_type(full_text)
    raw["occupancy_hint"]     = _detect_occupancy(full_text)

    # Sale History
    sale_section = re.search(
        r"Sale History from Public Records(.*?)(?:Mortgage History|Tax Information|\Z)",
        full_text, re.DOTALL | re.IGNORECASE
    )
    if sale_section:
        sh = sale_section.group(1)
        # Format 1: compact — seller+doc# on same line
        wd = re.search(
            r"(\d{1,2}/\d{1,2}/\d{2,4})\s+"
            r"(?:Y\s+)?([A-Za-z][A-Za-z &.]+?)\s*\n\s*"
            r"([A-Za-z][A-Za-z &.]+?)\s+(\d{4,6})\s*\n\s*"
            r"(Warranty Deed|Special Warranty Deed)",
            sh
        )
        if wd:
            raw["grantee_raw"]     = _clean(wd.group(2))
            raw["grantor_raw"]     = _clean(wd.group(3))
            raw["recorded_number"] = wd.group(4)
        else:
            # Format 2: expanded — each field on its own line, volume-page doc# (e.g. 8178-232)
            wd = re.search(
                r"(\d{1,2}/\d{1,2}/\d{2,4})\s*\n"
                r"\s*(?:Y\s*\n\s*)?([A-Za-z][A-Za-z &.,]+?)\s*\n"
                r"(?:\s*[A-Za-z][A-Za-z &.,]+?\s*\n\s*)?"
                r"\s*([A-Za-z][A-Za-z &.,]+?)\s*\n"
                r"\s*(\d{3,6}[-/]\d{1,6}|\d{4,6})\s*\n"
                r"\s*(Warranty Deed|Special Warranty Deed)",
                sh
            )
            if wd:
                raw["grantee_raw"]     = _clean(wd.group(2))
                raw["grantor_raw"]     = _clean(wd.group(3))
                raw["recorded_number"] = wd.group(4)
            else:
                wd2 = re.search(
                    r"(\d{3,6}[-/]\d{1,6}|\d{4,6})\s*\n?\s*(Warranty Deed|Special Warranty Deed)", sh
                )
                if wd2:
                    raw["recorded_number"] = wd2.group(1)

    return raw


def parse_mls_pdf(pdf_path):
    """
    Parse MLS PDF that may contain one or multiple listings.
    Returns list of raw dicts, each with _page_start and _page_end.
    """
    pdf        = fitz.open(pdf_path)
    pages_text = [page.get_text() for page in pdf]
    pdf.close()

    listing_starts = []
    for i, text in enumerate(pages_text):
        if (re.search(r"MLS#:\s*\d+", text) and
                re.search(r"[^\n]+,\s*[^,\n]+,\s*Texas\s+\d{5}", text, re.IGNORECASE)):
            listing_starts.append(i)

    if not listing_starts:
        listing_starts = [0]

    results = []
    for i, start in enumerate(listing_starts):
        end       = listing_starts[i + 1] if i + 1 < len(listing_starts) else len(pages_text)
        full_text = "\n".join(pages_text[start:end])
        p1        = pages_text[start]
        raw       = _parse_listing_text(full_text, p1)
        raw["_page_start"] = start
        raw["_page_end"]   = end
        results.append(raw)

    return results

# ============================================================
# CAD PDF PARSER
# ============================================================

def parse_cad_pdf(pdf_path):
    """
    Extract fields from a County Appraisal District PDF.
    Returns a dict that supplements MLS data (MLS values take priority on merge).
    """
    pdf        = fitz.open(pdf_path)
    pages_text = [page.get_text() for page in pdf]
    pdf.close()
    full = "\n".join(pages_text)

    raw = {}

    cad_m = re.search(r"^([A-Za-z][A-Za-z ]+?\s*CAD)\s+Property Search", full,
                      re.MULTILINE | re.IGNORECASE)
    raw["cad_source_name"] = _clean(cad_m.group(1)) if cad_m else "CAD"

    raw["cad_id"]  = _first(r"Property ID:\s*(\d+)", full)
    raw["geo_id"]  = _first(r"Geographic ID:\s*([\d.]+)", full)

    situs = _first(r"Situs Address:\s*([^\n]+)", full)
    if situs:
        sm = re.match(r"(.+?)\s+([A-Z][A-Z ]+),\s*TX\s+(\d{5})", situs.strip())
        if sm:
            raw["street"] = sm.group(1).title().strip()
            raw["city"]   = sm.group(2).title().strip()
            raw["state"]  = "Texas"

    raw["legal_desc"] = _first(r"Legal Description:\s*([^\n]+)", full)
    raw["owner_name"] = _first(r"\bName:\s*([A-Z][A-Z ]+[A-Z])\s*\n", full)

    # County from entity table e.g. "KC  KAUFMAN COUNTY  N/A"
    county_m = re.search(r"[A-Z]{2,4}\s+([A-Z]+ COUNTY)\s+N/A", full)
    if county_m:
        raw["county"] = county_m.group(1).title().replace(" County", "").strip()

    # Land acres and sqft from Property Land table
    land_m = re.search(r"\d+\w+\s+[A-Z][A-Z ]+\s+([\d.]+)\s+([\d,]+\.\d+)", full)
    if land_m:
        raw["acres"]  = land_m.group(1)
        raw["lot_sf"] = land_m.group(2).replace(",", "").split(".")[0]

    # Building SF from COMMERCIAL MAIN row
    bldg_m = re.search(r"COMMERCIAL MAIN\s+[\w*]+\s+(\d{4})\s+(\d+)", full, re.IGNORECASE)
    if bldg_m:
        raw["year_built"] = bldg_m.group(1)
        raw["bldg_sf"]    = bldg_m.group(2)

    # Also try RESIDENTIAL or other building types
    if not raw.get("year_built"):
        yr_m = re.search(r"Year Built[:\s]+(\d{4})", full, re.IGNORECASE)
        if yr_m:
            raw["year_built"] = yr_m.group(1)

    return raw


def merge_mls_cad(mls_raw, cad_raw):
    """MLS values win; CAD fills any fields MLS left blank."""
    merged = dict(mls_raw)
    for key, val in cad_raw.items():
        if val and not merged.get(key):
            merged[key] = val
    for key in ("cad_source_name", "bldg_sf", "year_built", "geo_id", "owner_name"):
        if cad_raw.get(key):
            merged[key] = cad_raw[key]
    return merged

# ============================================================
# FIELD DERIVATION
# ============================================================

# All 38 fields in display order (Excel uses all; Word skips blanks)
FIELD_ORDER = [
    "Comparable Sale",
    "Property Type",
    "CAD ID:",
    "Street Address:",
    "City:",
    "State:",
    "County:",
    "Land Size (SF):",
    "Land Size (Acres):",
    "Gross Building Area (SF):",
    "Net Rentable Area (SF)",
    "Rentable Unit Number",
    "Average Unit Size",
    "Unit Mix",
    "Project Amenities",
    "Unit Amenities",
    "Finish-out Percentage",
    "Land to Building Ratio:",
    "Year Built/ Renovated",
    "Date of Sale:",
    "Sales Price:",
    "Unit Price ($/Net SF):",
    "Unit Price ($/Gross SF):",
    "Unit Price ($/Unit)",
    "Cap Rate:",
    "NOI ($/Net SF)",
    "NOI ($/Unit)",
    "Occupancy Rate at Time of Sale:",
    "Data Source:",
    "Recorded  Number:",
    "Grantor:",
    "Grantee:",
    "Terms and Conditions:",
    "Time on Market:",
    "Property Rights Conveyed:",
    "Transactional Status",
    "Original Listing Price:",
    "Sale-to-List Ratio:",
    "Additional Comments:",
]

def derive_fields(raw, sale_number=1):
    d = {}

    d["Comparable Sale"] = f"Sale No. {sale_number}"

    # Property Type — detect from description, fall back to full-text hint
    desc = raw.get("description", "")
    d["Property Type"] = (_detect_property_type(desc)
                          or raw.get("property_type_hint", ""))

    d["CAD ID:"]         = raw.get("cad_id", "")
    d["Street Address:"] = raw.get("street", "")
    d["City:"]           = raw.get("city", "")
    d["State:"]          = raw.get("state", "Texas")
    d["County:"]         = raw.get("county", "")

    lot_sf   = raw.get("lot_sf", "")
    bldg_sf  = raw.get("bldg_sf", "")
    d["Land Size (SF):"]            = lot_sf
    d["Gross Building Area (SF):"]  = bldg_sf
    d["Net Rentable Area (SF)"]     = bldg_sf   # same as gross for single-tenant; analyst adjusts

    # Blanks — analyst fills for multi-unit or income properties
    d["Rentable Unit Number"] = ""
    d["Average Unit Size"]    = ""
    d["Unit Mix"]             = ""
    d["Project Amenities"]    = ""
    d["Unit Amenities"]       = ""

    # Finish-out percentage
    occupancy = (_detect_occupancy(desc) or raw.get("occupancy_hint", ""))
    d["Finish-out Percentage"] = "100%" if occupancy == "100%" else ""

    d["Year Built/ Renovated"] = raw.get("year_built", "")

    close_raw = raw.get("close_price", "")
    list_raw  = raw.get("list_price", "")
    d["Sales Price:"]            = _money_fmt(close_raw)
    d["Date of Sale:"]           = raw.get("close_date", "")
    d["Original Listing Price:"] = _money_fmt(list_raw)

    # Calculated fields
    if lot_sf and bldg_sf:
        try:
            ratio = float(lot_sf) / float(bldg_sf)
            d["Land to Building Ratio:"] = f"{ratio:.2f}"
        except (ValueError, ZeroDivisionError):
            d["Land to Building Ratio:"] = ""
    else:
        d["Land to Building Ratio:"] = ""

    if lot_sf:
        try:
            d["Land Size (Acres):"] = f"{float(lot_sf) / 43560:.3f}"
        except ValueError:
            d["Land Size (Acres):"] = ""
    else:
        d["Land Size (Acres):"] = ""

    if close_raw and bldg_sf:
        try:
            price = float(close_raw)
            sf    = float(bldg_sf)
            d["Unit Price ($/Net SF):"]   = f"${price / sf:,.2f}"
            d["Unit Price ($/Gross SF):"] = f"${price / sf:,.2f}"
        except (ValueError, ZeroDivisionError):
            d["Unit Price ($/Net SF):"]   = ""
            d["Unit Price ($/Gross SF):"] = ""
    else:
        d["Unit Price ($/Net SF):"]   = ""
        d["Unit Price ($/Gross SF):"] = ""

    d["Unit Price ($/Unit)"] = ""
    d["Cap Rate:"]           = ""
    d["NOI ($/Net SF)"]      = ""
    d["NOI ($/Unit)"]        = ""

    d["Occupancy Rate at Time of Sale:"] = (_detect_occupancy(desc)
                                             or raw.get("occupancy_hint", ""))

    mls_name   = raw.get("mls_name", "MLS")
    mls_number = raw.get("mls_number", "")
    cad_source = raw.get("cad_source_name", "")
    secondary  = cad_source if cad_source else "Texasfile.com"
    d["Data Source:"] = f"{mls_name} {mls_number}, {secondary}" if mls_number else ""

    d["Recorded  Number:"]        = raw.get("recorded_number", "")
    d["Grantor:"]                 = _flip_name(raw.get("grantor_raw", ""))
    d["Grantee:"]                 = _flip_name(raw.get("grantee_raw", ""))
    d["Terms and Conditions:"]    = _format_terms(raw.get("financing", ""))
    d["Time on Market:"]          = raw.get("dom", "")
    d["Property Rights Conveyed:"] = "Fee Simple"
    d["Transactional Status"]     = "Sold"

    if close_raw and list_raw:
        try:
            ratio = float(close_raw) / float(list_raw)
            d["Sale-to-List Ratio:"] = f"{ratio:.0%}"
        except (ValueError, ZeroDivisionError):
            d["Sale-to-List Ratio:"] = ""
    else:
        d["Sale-to-List Ratio:"] = ""

    # Additional comments from MLS description (first 300 chars as starter)
    d["Additional Comments:"] = desc[:300].strip() if desc else ""

    return d

# ============================================================
# EXCEL GENERATION
# ============================================================

def _border(style="thin"):
    s = Side(border_style=style, color="000000")
    return Border(left=s, right=s, top=s, bottom=s)

def generate_excel(all_comps, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    ws.column_dimensions["A"].width = 2
    ws.column_dimensions["B"].width = 32
    for col in ["C", "D", "E", "F"]:
        ws.column_dimensions[col].width = 24

    hdr_fill   = PatternFill("solid", fgColor="4472C4")
    label_fill = PatternFill("solid", fgColor="D9E1F2")
    alt_fill   = PatternFill("solid", fgColor="F2F2F2")
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    hdr_font   = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    label_font = Font(name="Calibri", size=10)
    val_font   = Font(name="Calibri", size=10)
    center_al  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_al    = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    # Row 1 — header
    ws.cell(row=1, column=2).value     = "Field"
    ws.cell(row=1, column=2).font      = hdr_font
    ws.cell(row=1, column=2).fill      = hdr_fill
    ws.cell(row=1, column=2).alignment = center_al
    ws.cell(row=1, column=2).border    = _border()

    for col_idx, comp in enumerate(all_comps, start=3):
        cell = ws.cell(row=1, column=col_idx)
        cell.value     = comp.get("Comparable Sale", f"Sale No. {col_idx - 2}")
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = center_al
        cell.border    = _border()

    # Data rows
    formula_fields = {
        "Land Size (Acres):":       lambda col: f"=ROUND({col}8/43560,3)",
        "Land to Building Ratio:":  lambda col: f"={col}8/{col}10",
        "Unit Price ($/Net SF):":   lambda col: f"={col}21/{col}11",
        "Unit Price ($/Gross SF):": lambda col: f"={col}21/{col}10",
        "Sale-to-List Ratio:":      lambda col: f"={col}21/{col}37",
    }
    col_letters = ["C", "D", "E", "F", "G", "H"]

    for row_idx, field in enumerate(FIELD_ORDER, start=2):
        use_alt    = (row_idx % 2 == 0)
        row_fill   = alt_fill if use_alt else white_fill

        label_cell = ws.cell(row=row_idx, column=2)
        label_cell.value     = field.rstrip(":").rstrip()
        label_cell.font      = label_font
        label_cell.fill      = label_fill
        label_cell.alignment = left_al
        label_cell.border    = _border()

        for ci, comp in enumerate(all_comps):
            col_letter = col_letters[ci] if ci < len(col_letters) else chr(ord("C") + ci)
            val_cell = ws.cell(row=row_idx, column=ci + 3)
            val_cell.font      = val_font
            val_cell.fill      = row_fill
            val_cell.alignment = left_al
            val_cell.border    = _border()

            # Use formula if this is a calculated field
            if field in formula_fields and comp.get(field, "") != "":
                val_cell.value = formula_fields[field](col_letter)
            else:
                val = comp.get(field, "")
                # Write numeric where appropriate
                for numeric_field in ("Land Size (SF):", "Gross Building Area (SF):",
                                      "Net Rentable Area (SF)", "Time on Market:"):
                    if field == numeric_field and val:
                        try:
                            val_cell.value = int(val)
                            break
                        except (ValueError, TypeError):
                            pass
                else:
                    val_cell.value = val

    ws.freeze_panes = "B2"
    wb.save(output_path)
    print(f"  Excel saved: {output_path}")

# ============================================================
# PHOTO EXTRACTION
# ============================================================

def extract_photos(pdf_path, tmp_dir, page_start=0, page_end=None, min_width=500):
    pdf      = fitz.open(pdf_path)
    page_end = page_end if page_end is not None else len(pdf)
    paths    = []

    for pg_num in range(page_start, page_end):
        page = pdf[pg_num]
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            pix  = fitz.Pixmap(pdf, xref)
            if pix.width < min_width:
                continue
            if pix.colorspace and pix.colorspace.n > 3:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img_path = os.path.join(tmp_dir, f"photo_{pg_num}_{xref}.png")
            pix.save(img_path)
            paths.append(img_path)

    pdf.close()
    return paths

# ============================================================
# WORD DOC GENERATION
# ============================================================

def _add_data_table(doc, comp):
    """Add data table showing only non-blank fields."""
    visible = [(f, comp.get(f, "")) for f in FIELD_ORDER if comp.get(f, "")]
    if not visible:
        return

    table = doc.add_table(rows=len(visible), cols=2)
    table.style = "Table Grid"

    for row_idx, (field, value) in enumerate(visible):
        row = table.rows[row_idx]

        label_cell = row.cells[0]
        label_cell.text = field.rstrip(":").rstrip()
        lp = label_cell.paragraphs[0]
        if lp.runs:
            lp.runs[0].font.bold = True
            lp.runs[0].font.size = Pt(9)
            lp.runs[0].font.name = "Calibri"
        label_cell.width = Inches(2.5)

        val_cell = row.cells[1]
        val_cell.text = str(value)
        vp = val_cell.paragraphs[0]
        if vp.runs:
            vp.runs[0].font.size = Pt(9)
            vp.runs[0].font.name = "Calibri"
        val_cell.width = Inches(3.8)

    return table

def _add_photo_grid(doc, photo_paths, img_width_inches=3.0):
    if not photo_paths:
        return
    n_rows = (len(photo_paths) + 1) // 2
    table  = doc.add_table(rows=n_rows, cols=2)
    table.style = "Table Grid"

    for i, photo_path in enumerate(photo_paths):
        cell = table.rows[i // 2].cells[i % 2]
        para = cell.paragraphs[0]
        run  = para.add_run()
        run.add_picture(photo_path, width=Inches(img_width_inches))

def generate_docx(all_comps_with_pdfs, output_path):
    doc     = Document()
    tmp_dir = tempfile.mkdtemp()

    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    try:
        for sale_idx, (comp, pdf_path, pg_start, pg_end) in enumerate(all_comps_with_pdfs, start=1):
            heading = doc.add_paragraph(style="Normal")
            run = heading.add_run(f"Improved Comparable Sale No. {sale_idx}")
            run.bold      = True
            run.font.size = Pt(12)
            run.font.name = "Calibri"

            _add_data_table(doc, comp)
            doc.add_paragraph()

            photos = extract_photos(pdf_path, tmp_dir, page_start=pg_start, page_end=pg_end)
            if photos:
                _add_photo_grid(doc, photos)
                doc.add_paragraph()

    finally:
        for f in os.listdir(tmp_dir):
            os.remove(os.path.join(tmp_dir, f))
        os.rmdir(tmp_dir)

    doc.save(output_path)
    print(f"  Word doc saved: {output_path}")

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate ImprovedComp Excel + Word doc")
    parser.add_argument("--name",      help="Output name (e.g. 'Terrell')")
    parser.add_argument("--mls",       action="append", metavar="PDF", help="MLS PDF path")
    parser.add_argument("--cad",       action="append", metavar="PDF", help="CAD PDF path (optional)")
    parser.add_argument("--texasfile", action="append", metavar="PNG", help="TexasFile PNG (optional)")
    args = parser.parse_args()

    if args.name:
        OUTPUT_NAME = args.name
    if args.mls:
        MLS_PDF_PATHS = args.mls
    if args.cad:
        CAD_PDF_PATHS = args.cad
    if args.texasfile:
        TEXASFILE_PNG_PATHS = args.texasfile

    if not MLS_PDF_PATHS:
        print("ERROR: No MLS PDFs specified. Edit CONFIG or use --mls argument.")
        raise SystemExit(1)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Parse CAD PDFs into lookup by Property ID
    cad_by_id = {}
    cad_list  = []
    if CAD_PDF_PATHS:
        print("Parsing CAD PDFs...")
        for cad_path in CAD_PDF_PATHS:
            cad_raw = parse_cad_pdf(cad_path)
            print(f"  {os.path.basename(cad_path)}")
            print(f"    Property ID: {cad_raw.get('cad_id','?')}  "
                  f"County: {cad_raw.get('county','?')}  "
                  f"Bldg SF: {cad_raw.get('bldg_sf','?')}  "
                  f"Year Built: {cad_raw.get('year_built','?')}")
            cad_list.append(cad_raw)
            if cad_raw.get("cad_id"):
                cad_by_id[cad_raw["cad_id"]] = cad_raw

    print("\nParsing MLS PDFs...")
    all_comps           = []
    all_comps_with_pdfs = []
    sale_number         = 0
    listing_index       = 0

    for pdf_path in MLS_PDF_PATHS:
        listings = parse_mls_pdf(pdf_path)
        print(f"  {os.path.basename(pdf_path)}  ({len(listings)} listing(s))")

        for raw in listings:
            sale_number   += 1
            listing_index += 1

            cad_raw = cad_by_id.get(raw.get("cad_id", ""))
            if cad_raw is None and listing_index <= len(cad_list):
                cad_raw = cad_list[listing_index - 1]
            if cad_raw:
                raw = merge_mls_cad(raw, cad_raw)
                print(f"    Merged: {cad_raw.get('cad_source_name','CAD')} "
                      f"(ID {cad_raw.get('cad_id','?')})")

            comp = derive_fields(raw, sale_number=sale_number)
            all_comps.append(comp)
            all_comps_with_pdfs.append((comp, pdf_path, raw["_page_start"], raw["_page_end"]))

            print(f"    Sale {sale_number}: {raw.get('street','?')}, {raw.get('city','?')}")
            print(f"      MLS#: {raw.get('mls_number','?')}  CAD ID: {raw.get('cad_id','?')}")
            print(f"      Lot SF: {raw.get('lot_sf','?')}  Bldg SF: {raw.get('bldg_sf','?')}  "
                  f"Yr Built: {raw.get('year_built','?')}")
            print(f"      Price: {raw.get('close_price','?')}  OLP: {raw.get('list_price','?')}  "
                  f"DOM: {raw.get('dom','?')}")
            print(f"      Type detected: {comp.get('Property Type','(blank)')}")
            print(f"      Grantor: {raw.get('grantor_raw','?')}  "
                  f"Grantee: {raw.get('grantee_raw','?')}")
            print(f"      RecNum: {raw.get('recorded_number','?')}  "
                  f"Financing: {raw.get('financing','?')}")

    print("\nGenerating Excel...")
    xlsx_path = os.path.join(OUTPUT_DIR, f"ImprovedComp_{OUTPUT_NAME}.xlsx")
    generate_excel(all_comps, xlsx_path)

    print("\nGenerating Word doc...")
    docx_path = os.path.join(OUTPUT_DIR, f"ImprovedComp_{OUTPUT_NAME}.docx")
    generate_docx(all_comps_with_pdfs, docx_path)

    print(f"\nDone. Output in: {OUTPUT_DIR}")
