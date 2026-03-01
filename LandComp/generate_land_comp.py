"""
Land Comparable Generator — generate_land_comp.py

Parses one or more MLS PDFs to generate:
  - LandComp_[name].xlsx   (comparison spreadsheet)
  - LandComp_[name].docx   (Word doc with tables + MLS photos)

Usage — edit CONFIG below, then run:
    cd c:\\AppraisalWorkspace\\LandComp
    python generate_land_comp.py

Usage — command-line arguments (overrides CONFIG):
    python generate_land_comp.py --name "Corsicana" --mls "path/to/MLS1.pdf" --texasfile "path/to/tf.png"

    For multiple comparables, repeat --mls and --texasfile:
    python generate_land_comp.py --name "Austin" --mls "MLS1.pdf" --mls "MLS2.pdf"
"""

import os
import re
import sys
import argparse
import tempfile
import fitz  # PyMuPDF
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONFIG — Update these for each new report
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCES_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), "Sources")

MLS_PDF_PATHS = [
    os.path.join(SOURCES_DIR, "MLS", "MLS-Corsicana, Texas, 75110.pdf"),
    # Add more PDFs here for additional comparables:
    # os.path.join(SOURCES_DIR, "MLS", "MLS-AnotherProperty.pdf"),
]

# TexasFile PNGs are optional — used for Data Source field notation only
TEXASFILE_PNG_PATHS = [
    os.path.join(SOURCES_DIR, "TexasFile", "texasfile-navarroSearch.png"),
]

OUTPUT_DIR  = os.path.join(SCRIPT_DIR, "Output")
OUTPUT_NAME = "Corsicana"  # used in output filenames

# ============================================================
# HELPERS
# ============================================================

def _clean(s):
    """Strip whitespace and non-breaking spaces."""
    if not s:
        return ""
    return s.replace("\xa0", " ").replace("\u00b0", "").strip()

def _first(pattern, text, group=1, flags=re.IGNORECASE | re.MULTILINE):
    m = re.search(pattern, text, flags)
    return _clean(m.group(group)) if m else ""

def _money_fmt(raw_str):
    """'51500' or '51,500' → '$51,500'"""
    s = str(raw_str).replace(",", "").replace("$", "").strip()
    if not s:
        return ""
    try:
        v = float(s)
        return f"${v:,.0f}" if v == int(v) else f"${v:,.2f}"
    except ValueError:
        return ""

def _flip_name(name):
    """
    Reformat a name from MLS format to display format.
    'Perez Fernando'        → 'Fernando Perez'
    'Wallen Joshua & Brandi'→ 'Joshua and Brandi Wallen'
    """
    name = _clean(name)
    if not name:
        return ""
    if "&" in name:
        parts = [p.strip() for p in name.split("&")]
        first_parts = parts[0].split()
        if len(first_parts) >= 2:
            last = first_parts[0]
            first1 = " ".join(first_parts[1:])
            first2 = parts[1].strip()
            return f"{first1} and {first2} {last}"
        return name
    parts = name.split()
    if len(parts) >= 2:
        return " ".join(parts[1:]) + " " + parts[0]
    return name

# ============================================================
# MLS PDF PARSER
# ============================================================

ZONING_MAP = {
    "general retail":      "GR (General Retail)",
    "commercial":          "C (Commercial)",
    "light industrial":    "LI (Light Industrial)",
    "heavy industrial":    "HI (Heavy Industrial)",
    "residential":         "R (Residential)",
    "agricultural":        "A (Agricultural)",
    "mixed use":           "MU (Mixed Use)",
    "office":              "O (Office)",
    "planned development": "PD (Planned Development)",
}

def _format_zoning(raw):
    key = raw.lower().strip()
    for k, v in ZONING_MAP.items():
        if k in key:
            return v
    return raw.strip()

def _format_utilities(raw):
    r = raw.lower()
    if "all" in r:
        return "All available"
    # City/commercial lots listing any utility (electricity, water, sewer, etc.)
    # typically have all utilities at the street; default to "All available"
    if r and any(w in r for w in ["electric", "water", "sewer", "gas", "avail"]):
        return "All available"
    if r:
        return "Available"
    return ""

def _format_terms(financing):
    f = financing.lower().strip()
    if "cash" in f:
        return "Cash or conventional financing"
    if "conventional" in f:
        return "Conventional financing"
    if "fha" in f:
        return "FHA financing"
    if "va" in f:
        return "VA financing"
    return financing.strip().title() if financing.strip() else ""

def _parse_listing_text(full_text, p1):
    """
    Extract raw field values from a single listing's text block.
    p1 = first page text of the listing (used for address extraction).
    Returns a dict of raw string values.
    """
    raw = {}

    # ── Address (first line: "Street, City, Texas XXXXX") ───────────────────
    addr_m = re.search(
        r"[^\w\n]*([^\n,]+),\s*([^,\n]+),\s*Texas\s+(\d{5})",
        p1, re.IGNORECASE
    )
    if addr_m:
        raw["street"] = _clean(addr_m.group(1))
        raw["city"]   = _clean(addr_m.group(2))
        raw["state"]  = "Texas"

    # ── MLS number ───────────────────────────────────────────────────────────
    raw["mls_number"] = _first(r"MLS#:\s*(\d+)", full_text)

    # ── County (value sometimes on next line) ────────────────────────────────
    raw["county"] = _first(r"County:\s*\n?\s*([A-Za-z][A-Za-z ]*?)(?:\n|Lake Name|$)", full_text)

    # ── CAD / Parcel ID ──────────────────────────────────────────────────────
    raw["cad_id"] = _first(r"Parcel ID:\s*\n?\s*(\d+)", full_text)

    # ── Lot SF ───────────────────────────────────────────────────────────────
    raw["lot_sf"] = _first(r"Lot SqFt:\s*([\d,]+)", full_text).replace(",", "")

    # ── Acres ────────────────────────────────────────────────────────────────
    raw["acres"] = _first(r"(?:^|[\s\n])Acres:\s*\n?\s*([\d.]+)", full_text)

    # ── Close Price ──────────────────────────────────────────────────────────
    raw["close_price"] = _first(r"Close Price:\s*\$([\d,]+)", full_text).replace(",", "")

    # ── Close Date ───────────────────────────────────────────────────────────
    raw["close_date"] = _first(r"Closed Date:\s*(\d{1,2}/\d{1,2}/\d{4})", full_text)

    # ── Original Listing Price (OLP preferred, fall back to LP) ─────────────
    lp = _first(r"OLP:\s*\$([\d,]+)", full_text).replace(",", "")
    if not lp:
        lp = _first(r"(?:^|\n)LP:\s*\$?\s*\n?\s*\$?([\d,]+)", full_text).replace(",", "")
    raw["list_price"] = lp

    # ── Zoning ───────────────────────────────────────────────────────────────
    raw["zoning"] = _first(r"Zoning:\s*([^\n]+)", full_text)

    # ── DOM / CDOM ───────────────────────────────────────────────────────────
    cdom = _first(r"\bCDOM:\s*(\d+)", full_text)
    dom  = _first(r"\bDOM:\s*(\d+)", full_text)
    raw["dom"] = cdom or dom

    # ── Flood Zone ───────────────────────────────────────────────────────────
    raw["flood_zone"] = _first(r"Flood Zone Code:\s*(\w+)", full_text)

    # ── Utilities ────────────────────────────────────────────────────────────
    raw["utilities_raw"] = _first(r"Street/Utilities:\s*([^\n]+)", full_text)

    # ── Buyer Financing ──────────────────────────────────────────────────────
    raw["financing"] = _first(r"Buyer Financing:\s*([^\n]+)", full_text)

    # ── MLS name (from copyright line) ───────────────────────────────────────
    if "NTREIS" in full_text:
        raw["mls_name"] = "North Texas MLS"
    elif "HAR" in full_text:
        raw["mls_name"] = "Houston MLS"
    else:
        raw["mls_name"] = "MLS"

    # ── Sale History — most recent Warranty Deed ─────────────────────────────
    sale_section = re.search(
        r"Sale History from Public Records(.*?)(?:Mortgage History|Tax Information|\Z)",
        full_text, re.DOTALL | re.IGNORECASE
    )
    if sale_section:
        sh = sale_section.group(1)
        wd = re.search(
            r"(\d{1,2}/\d{1,2}/\d{2,4})\s+"
            r"(?:Y\s+)?([A-Za-z][A-Za-z &.]+?)\s*\n\s*"
            r"([A-Za-z][A-Za-z &.]+?)\s+(\d{4,6})\s*\n\s*"
            r"(Warranty Deed|Special Warranty Deed)",
            sh
        )
        if wd:
            raw["grantee_raw"]     = _clean(wd.group(2))
            raw["grantor_raw"]     = _clean(wd.group(3))
            raw["recorded_number"] = wd.group(4)
            raw["deed_type"]       = wd.group(5)
        else:
            wd2 = re.search(r"(\d{4,6})\s*\n?\s*(Warranty Deed|Special Warranty Deed)", sh)
            if wd2:
                raw["recorded_number"] = wd2.group(1)
                raw["deed_type"]       = wd2.group(2)

    return raw


def parse_mls_pdf(pdf_path):
    """
    Parse a MLS PDF that may contain one OR multiple listings.

    Each listing's first page is identified by having both an MLS# and a
    Texas address line.  All pages between one listing header and the next
    are treated as belonging to that listing (for text extraction and photos).

    Returns a list of dicts — one per listing found.
    Each dict has all raw field values plus:
        '_page_start'  (int) — 0-based index of first PDF page for this listing
        '_page_end'    (int) — 0-based index (exclusive) of last PDF page
    """
    pdf        = fitz.open(pdf_path)
    pages_text = [page.get_text() for page in pdf]
    pdf.close()

    # ── Find page indices where a new listing starts ──────────────────────
    listing_starts = []
    for i, text in enumerate(pages_text):
        has_mls  = bool(re.search(r"MLS#:\s*\d+", text))
        has_addr = bool(re.search(r"[^\n]+,\s*[^,\n]+,\s*Texas\s+\d{5}", text, re.IGNORECASE))
        if has_mls and has_addr:
            listing_starts.append(i)

    if not listing_starts:
        listing_starts = [0]  # fall back: treat whole PDF as one listing

    # ── Parse each listing's page range ──────────────────────────────────
    results = []
    for i, start in enumerate(listing_starts):
        end       = listing_starts[i + 1] if i + 1 < len(listing_starts) else len(pages_text)
        full_text = "\n".join(pages_text[start:end])
        p1        = pages_text[start]
        raw       = _parse_listing_text(full_text, p1)
        raw["_page_start"] = start
        raw["_page_end"]   = end
        results.append(raw)

    return results

# ============================================================
# FIELD DERIVATION
# ============================================================

# Ordered list of (field_label, key_in_derived)
FIELD_ORDER = [
    "Comparable Sale",
    "Property Type",
    "CAD ID:",
    "Street Address:",
    "City:",
    "State:",
    "County:",
    "Land Size (SF):",
    "Land Size (Acres):",
    "Sales Price:",
    "Date of Sale:",
    "Unit Price (S/SF):",
    "Unit Price (S/Acre):",
    "Zoning:",
    "Configuration:",
    "Topography:",
    "Min Topo (Ft)",
    "Max Topo (Ft)",
    "Topo % Change",
    "Access:",
    "Utilities:",
    "Flood Zone:",
    "Data Source:",
    "Recorded Number:",
    "Grantor:",
    "Grantee:",
    "Terms and Conditions:",
    "Days on Market:",
    "Property Rights Conveyed:",
    "Transactional Status",
    "Original Listing Price:",
    "Sale-to-List Ratio:",
    "Additional Comments:",
]

def derive_fields(raw, sale_number=1):
    """Convert raw parsed dict → display-ready dict keyed by FIELD_ORDER labels."""
    d = {}

    d["Comparable Sale"]   = f"Sale No. {sale_number}"
    d["Property Type"]     = "Land Site"
    d["CAD ID:"]           = raw.get("cad_id", "")
    d["Street Address:"]   = raw.get("street", "")
    d["City:"]             = raw.get("city", "")
    d["State:"]            = raw.get("state", "Texas")
    d["County:"]           = raw.get("county", "")

    lot_sf = raw.get("lot_sf", "")
    acres  = raw.get("acres", "")
    d["Land Size (SF):"]    = lot_sf
    d["Land Size (Acres):"] = acres

    close_raw = raw.get("close_price", "")
    list_raw  = raw.get("list_price", "")
    d["Sales Price:"]           = _money_fmt(close_raw)
    d["Date of Sale:"]          = raw.get("close_date", "")
    d["Original Listing Price:"] = _money_fmt(list_raw)

    # Unit prices — calculated from SF for precision
    if close_raw and lot_sf:
        try:
            price  = float(close_raw)
            sf_val = float(lot_sf)
            d["Unit Price (S/SF):"]   = f"${price / sf_val:,.2f}"
            d["Unit Price (S/Acre):"] = f"${price * 43560 / sf_val:,.0f}"
        except (ValueError, ZeroDivisionError):
            d["Unit Price (S/SF):"]   = ""
            d["Unit Price (S/Acre):"] = ""
    else:
        d["Unit Price (S/SF):"]   = ""
        d["Unit Price (S/Acre):"] = ""

    # Sale-to-List Ratio
    if close_raw and list_raw:
        try:
            ratio = float(close_raw) / float(list_raw)
            d["Sale-to-List Ratio:"] = f"{ratio:.0%}"
        except (ValueError, ZeroDivisionError):
            d["Sale-to-List Ratio:"] = ""
    else:
        d["Sale-to-List Ratio:"] = ""

    d["Zoning:"]        = _format_zoning(raw.get("zoning", ""))
    d["Configuration:"] = "Regular"
    d["Topography:"]    = ""          # blank — human fills from site visit
    d["Min Topo (Ft)"]  = ""          # blank — human fills from Google Maps
    d["Max Topo (Ft)"]  = ""          # blank — human fills from Google Maps
    d["Topo % Change"]  = ""          # blank — calculated by human
    d["Access:"]        = "Inside"    # default for non-corner city/town lots

    d["Utilities:"]  = _format_utilities(raw.get("utilities_raw", ""))
    d["Flood Zone:"] = raw.get("flood_zone", "")

    mls_name   = raw.get("mls_name", "MLS")
    mls_number = raw.get("mls_number", "")
    d["Data Source:"] = f"{mls_name} {mls_number}, Texasfile.com" if mls_number else ""

    d["Recorded Number:"] = raw.get("recorded_number", "")
    d["Grantor:"]         = _flip_name(raw.get("grantor_raw", ""))
    d["Grantee:"]         = _flip_name(raw.get("grantee_raw", ""))

    d["Terms and Conditions:"]    = _format_terms(raw.get("financing", ""))
    d["Days on Market:"]          = raw.get("dom", "")
    d["Property Rights Conveyed:"] = "Fee Simple"
    d["Transactional Status"]     = "Sold"
    d["Additional Comments:"]     = "None."

    return d

# ============================================================
# EXCEL GENERATION
# ============================================================

def _border(style="thin"):
    s = Side(border_style=style, color="000000")
    return Border(left=s, right=s, top=s, bottom=s)

def generate_excel(all_comps, output_path):
    """
    all_comps: list of derived field dicts (one per comparable)
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    # Column widths
    ws.column_dimensions["A"].width = 30
    for col_letter in ["B", "C", "D", "E", "F"]:
        ws.column_dimensions[col_letter].width = 22

    # Styles
    hdr_font    = Font(name="Calibri", bold=True, size=11)
    label_font  = Font(name="Calibri", size=10)
    value_font  = Font(name="Calibri", size=10)
    label_fill  = PatternFill("solid", fgColor="D9E1F2")   # light blue
    alt_fill    = PatternFill("solid", fgColor="F2F2F2")   # light grey
    center_al   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_al     = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    # Header row (row 1): field labels column + one column per comp
    ws.cell(row=1, column=1).value = "Field"
    ws.cell(row=1, column=1).font  = hdr_font
    ws.cell(row=1, column=1).fill  = PatternFill("solid", fgColor="4472C4")
    ws.cell(row=1, column=1).font  = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    ws.cell(row=1, column=1).alignment = center_al
    ws.cell(row=1, column=1).border    = _border()

    for col_idx, comp in enumerate(all_comps, start=2):
        cell = ws.cell(row=1, column=col_idx)
        cell.value = comp.get("Comparable Sale", f"Sale No. {col_idx - 1}")
        cell.font  = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
        cell.fill  = PatternFill("solid", fgColor="4472C4")
        cell.alignment = center_al
        cell.border    = _border()

    # Data rows
    for row_idx, field in enumerate(FIELD_ORDER, start=2):
        use_alt = (row_idx % 2 == 0)
        row_fill = alt_fill if use_alt else PatternFill("solid", fgColor="FFFFFF")

        label_cell = ws.cell(row=row_idx, column=1)
        label_cell.value     = field.rstrip(":").rstrip()
        label_cell.font      = label_font
        label_cell.fill      = label_fill
        label_cell.alignment = left_al
        label_cell.border    = _border()

        for col_idx, comp in enumerate(all_comps, start=2):
            val_cell = ws.cell(row=row_idx, column=col_idx)
            val_cell.value     = comp.get(field, "")
            val_cell.font      = value_font
            val_cell.fill      = row_fill
            val_cell.alignment = left_al
            val_cell.border    = _border()

    # Freeze top row
    ws.freeze_panes = "A2"

    wb.save(output_path)
    print(f"  Excel saved: {output_path}")

# ============================================================
# PHOTO EXTRACTION
# ============================================================

def extract_photos(pdf_path, tmp_dir, page_start=0, page_end=None, min_width=500):
    """
    Extract large images from a listing's page range within a MLS PDF.
    Scans all pages in [page_start, page_end) and returns paths to
    saved PNGs for images wider than min_width pixels.
    """
    pdf      = fitz.open(pdf_path)
    page_end = page_end if page_end is not None else len(pdf)
    paths    = []

    for pg_num in range(page_start, page_end):
        page = pdf[pg_num]
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            pix  = fitz.Pixmap(pdf, xref)
            if pix.width < min_width:
                continue
            if pix.colorspace and pix.colorspace.n > 3:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img_path = os.path.join(tmp_dir, f"photo_{pg_num}_{xref}.png")
            pix.save(img_path)
            paths.append(img_path)

    pdf.close()
    return paths

# ============================================================
# WORD DOC GENERATION
# ============================================================

def _set_cell_border(cell):
    """Apply thin borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), "000000")
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(el)

def _add_data_table(doc, comp):
    """Add a 2-column data table for one comparable."""
    table = doc.add_table(rows=len(FIELD_ORDER), cols=2)
    table.style = "Table Grid"

    col_widths = [Inches(2.5), Inches(3.8)]

    for row_idx, field in enumerate(FIELD_ORDER):
        row = table.rows[row_idx]

        # Label cell
        label_cell = row.cells[0]
        label_cell.text = field.rstrip(":").rstrip()
        label_p = label_cell.paragraphs[0]
        label_p.runs[0].font.bold = True
        label_p.runs[0].font.size = Pt(9)
        label_p.runs[0].font.name = "Calibri"
        label_cell.width = col_widths[0]

        # Value cell
        val_cell = row.cells[1]
        val_cell.text = str(comp.get(field, ""))
        val_p = val_cell.paragraphs[0]
        if val_p.runs:
            val_p.runs[0].font.size = Pt(9)
            val_p.runs[0].font.name = "Calibri"
        val_cell.width = col_widths[1]

    return table

def _add_photo_grid(doc, photo_paths, img_width_inches=3.0):
    """Add a 2-column image grid table with MLS photos."""
    if not photo_paths:
        return

    n_rows = (len(photo_paths) + 1) // 2
    table  = doc.add_table(rows=n_rows, cols=2)
    table.style = "Table Grid"

    for i, photo_path in enumerate(photo_paths):
        row_idx = i // 2
        col_idx = i % 2
        cell = table.rows[row_idx].cells[col_idx]
        para = cell.paragraphs[0]
        run  = para.add_run()
        run.add_picture(photo_path, width=Inches(img_width_inches))

def generate_docx(all_comps_with_pdfs, output_path):
    """
    all_comps_with_pdfs: list of (derived_dict, mls_pdf_path, page_start, page_end) tuples
    """
    doc     = Document()
    tmp_dir = tempfile.mkdtemp()

    # Page margins: 1 inch
    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    try:
        for sale_idx, (comp, pdf_path, pg_start, pg_end) in enumerate(all_comps_with_pdfs, start=1):
            # Heading: "Land Sale Comp No. X"
            heading = doc.add_paragraph(style="Normal")
            run = heading.add_run(f"Land Sale Comp No. {sale_idx}")
            run.bold      = True
            run.font.size = Pt(12)
            run.font.name = "Calibri"

            # Data table
            _add_data_table(doc, comp)
            doc.add_paragraph()  # spacer

            # Photos — only from this listing's page range
            photos = extract_photos(pdf_path, tmp_dir, page_start=pg_start, page_end=pg_end)
            if photos:
                _add_photo_grid(doc, photos)
                doc.add_paragraph()  # spacer after photos

    finally:
        for f in os.listdir(tmp_dir):
            os.remove(os.path.join(tmp_dir, f))
        os.rmdir(tmp_dir)

    doc.save(output_path)
    print(f"  Word doc saved: {output_path}")

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    # ── CLI args override CONFIG values ──────────────────────────────────────
    parser = argparse.ArgumentParser(description="Generate LandComp Excel + Word doc from MLS PDFs")
    parser.add_argument("--name",      help="Output name (e.g. 'Corsicana')")
    parser.add_argument("--mls",       action="append", metavar="PDF",  help="MLS PDF path (repeat for multiple)")
    parser.add_argument("--texasfile", action="append", metavar="PNG",  help="TexasFile PNG path (optional, repeat to match)")
    args = parser.parse_args()

    if args.name:
        OUTPUT_NAME = args.name
    if args.mls:
        MLS_PDF_PATHS = args.mls
    if args.texasfile:
        TEXASFILE_PNG_PATHS = args.texasfile

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("Parsing MLS PDFs...")
    all_comps           = []
    all_comps_with_pdfs = []
    sale_number         = 0

    for pdf_path in MLS_PDF_PATHS:
        listings = parse_mls_pdf(pdf_path)   # returns list — 1 entry per listing found
        print(f"  {os.path.basename(pdf_path)}  ({len(listings)} listing(s) detected)")

        for raw in listings:
            sale_number += 1
            comp = derive_fields(raw, sale_number=sale_number)
            all_comps.append(comp)
            all_comps_with_pdfs.append((
                comp,
                pdf_path,
                raw["_page_start"],
                raw["_page_end"],
            ))

            # Debug: print extracted values
            print(f"    Sale {sale_number}: {raw.get('street','?')}, {raw.get('city','?')}")
            print(f"      MLS#:     {raw.get('mls_number','(not found)')}")
            print(f"      CAD ID:   {raw.get('cad_id','(not found)')}")
            print(f"      Lot SF:   {raw.get('lot_sf','(not found)')}  Acres: {raw.get('acres','(not found)')}")
            print(f"      Price:    {raw.get('close_price','(not found)')}  Date: {raw.get('close_date','(not found)')}")
            print(f"      OLP:      {raw.get('list_price','(not found)')}  DOM: {raw.get('dom','(not found)')}")
            print(f"      Zoning:   {raw.get('zoning','(not found)')}  Flood: {raw.get('flood_zone','(not found)')}")
            print(f"      Grantor:  {raw.get('grantor_raw','(not found)')}")
            print(f"      Grantee:  {raw.get('grantee_raw','(not found)')}")
            print(f"      RecNum:   {raw.get('recorded_number','(not found)')}  Financing: {raw.get('financing','(not found)')}")
            print(f"      Pages:    {raw['_page_start']+1}–{raw['_page_end']} of PDF")

    print("\nGenerating Excel...")
    xlsx_path = os.path.join(OUTPUT_DIR, f"LandComp_{OUTPUT_NAME}.xlsx")
    generate_excel(all_comps, xlsx_path)

    print("\nGenerating Word doc...")
    docx_path = os.path.join(OUTPUT_DIR, f"LandComp_{OUTPUT_NAME}.docx")
    generate_docx(all_comps_with_pdfs, docx_path)

    print(f"\nDone. Output in: {OUTPUT_DIR}")
