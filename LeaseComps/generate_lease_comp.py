"""
Lease Comparable Generator — generate_lease_comp.py
Parses a Collier lease comps PDF and generates:
  - LeaseComp_[name].xlsx   (comparison spreadsheet matching example format)
  - LeaseComp_[name].docx   (Word doc with data table + PDF page images)

Usage:
    cd c:\\AppraisalWorkspace\\LeaseComps
    python generate_lease_comp.py

    Or with CLI args:
    python generate_lease_comp.py --name "CStore" --pdf "path/to/lease_comps.pdf"

Source PDF format: Colliers International — lease comparable pages
  - Multiple comparables per page (2 per page typical)
  - Structured fields: Physical Info, Confirmation, Tenant/Lease table, Remarks

IMPORTANT — Document creation:
  The Word doc is built from a blank Document() — no template is loaded as base.
  This avoids orphaned image blobs that would bloat the file and crash Word on
  copy/paste. See RetailMarketReport/README.txt for detailed explanation.
"""

import os
import re
import io
import argparse
import fitz  # PyMuPDF
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ============================================================
# CONFIG
# ============================================================
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
SOURCES_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), "Sources")

PDF_PATH    = os.path.join(SOURCES_DIR, "Comps", "lease cstore comps - collier.pdf")
OUTPUT_DIR  = os.path.join(SCRIPT_DIR, "Output")
OUTPUT_NAME = "Collier_CStore"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Image rendering
RENDER_DPI     = 200
IMAGE_WIDTH_IN = 6.5
JPEG_QUALITY   = 85

# ============================================================
# HELPERS
# ============================================================

def _clean(s):
    if not s:
        return ""
    return s.replace("\xa0", " ").replace("\u00b0", "").strip()

def _money_fmt(raw_str):
    s = str(raw_str).replace(",", "").replace("$", "").strip()
    if not s:
        return ""
    try:
        v = float(s)
        return f"${v:,.2f}"
    except ValueError:
        return ""

# ============================================================
# COLLIER LEASE PDF PARSER
# ============================================================

def _split_comparables(page_text):
    """Split a page's text into individual comparable blocks."""
    # Split on "COMPARABLE N" headers
    parts = re.split(r"(?=COMPARABLE\s+\d+)", page_text)
    return [p for p in parts if re.match(r"COMPARABLE\s+\d+", p.strip())]

def _extract_field(text, label):
    """Extract value that follows a label on the next line."""
    pattern = re.escape(label) + r"\s*\n\s*(.+)"
    m = re.search(pattern, text)
    return _clean(m.group(1)) if m else ""

def parse_lease_comparable(comp_text):
    """Parse one lease comparable block into a raw dict."""
    raw = {}

    # Comparable number
    comp_m = re.search(r"COMPARABLE\s+(\d+)", comp_text)
    raw["comp_number"] = int(comp_m.group(1)) if comp_m else 0

    # Physical Information
    raw["name"] = _extract_field(comp_text, "Name")
    raw["address"] = _extract_field(comp_text, "Address")

    csz = _extract_field(comp_text, "City, State, Zip Code")
    if csz:
        # Handle formats like "Port Aransas, TX 78373" or "Corpus Christi, TX, TX 78415"
        csz_m = re.match(r"(.+?),\s*(?:TX,?\s*)?TX?\s*(\d{5})", csz)
        if csz_m:
            raw["city"] = csz_m.group(1).strip()
            raw["state"] = "Texas"
            raw["zip"] = csz_m.group(2)
        else:
            raw["city"] = csz.split(",")[0].strip()
            raw["state"] = "Texas"

    raw["msa"] = _extract_field(comp_text, "MSA")

    nra = _extract_field(comp_text, "Net Rentable Area (NRA)")
    raw["nra"] = nra.replace(",", "") if nra else ""

    raw["year_built"] = _extract_field(comp_text, "Year Built")
    raw["occupancy"] = _extract_field(comp_text, "Occupancy")

    site_size = _extract_field(comp_text, "Site Size")
    raw["site_sf"] = site_size.replace(",", "") if site_size else ""

    raw["site_coverage"] = _extract_field(comp_text, "Ste Coverage")
    if not raw["site_coverage"]:
        raw["site_coverage"] = _extract_field(comp_text, "Site Coverage")
    raw["construction"] = _extract_field(comp_text, "Construction")

    # Confirmation
    raw["company"] = _extract_field(comp_text, "Company")
    raw["source"] = _extract_field(comp_text, "Source")

    # Lease table — headers and data are each on their own line:
    #   TENANT NAME\nRATE TYPE\nSIZE\nSTART DATE\nTERM\nLEASE RATE\nADJ LEASE RATE
    #   Valero\nNNN\n7,425\nCurrent\n20\n$21.68\n$23.85
    lease_m = re.search(
        r"TENANT NAME\s*\nRATE TYPE\s*\nSIZE\s*\nSTART DATE\s*\nTERM\s*\nLEASE RATE\s*\nADJ LEASE RATE\s*\n"
        r"(.+?)\n"          # tenant name
        r"(.+?)\n"          # rate type
        r"([\d,]+)\n"      # size
        r"(.+?)\n"          # start date
        r"(\d+)\n"         # term (years)
        r"\$([\.\d]+)\n"   # lease rate
        r"\$([\.\d]+)",    # adj lease rate
        comp_text
    )
    if lease_m:
        raw["tenant_name"] = _clean(lease_m.group(1))
        raw["rate_type"] = _clean(lease_m.group(2))
        raw["lease_size"] = lease_m.group(3).replace(",", "")
        raw["start_date"] = _clean(lease_m.group(4))
        raw["term_years"] = lease_m.group(5)
        raw["lease_rate"] = lease_m.group(6)
        raw["adj_lease_rate"] = lease_m.group(7)

    # Remarks — text after the adjusted lease rate line (last $ value before next COMPARABLE or end)
    # The adj_lease_rate is the last data field; everything after it is the remark sentence.
    adj_rate = raw.get("adj_lease_rate", "")
    if adj_rate:
        # Find text after "$<adj_rate>\n" up to next COMPARABLE or end
        pattern = r"\$" + re.escape(adj_rate) + r"\s*\n(.*?)(?:COMPARABLE|\Z)"
        remarks_m = re.search(pattern, comp_text, re.DOTALL)
        if remarks_m:
            raw["remarks"] = _clean(remarks_m.group(1).replace("\n", " "))
        else:
            raw["remarks"] = ""
    else:
        raw["remarks"] = ""

    return raw


def parse_lease_pdf(pdf_path):
    """Parse a Collier lease comps PDF. Returns list of raw dicts with page indices."""
    pdf = fitz.open(pdf_path)
    results = []
    for i in range(pdf.page_count):
        text = pdf[i].get_text()
        comp_blocks = _split_comparables(text)
        for block in comp_blocks:
            raw = parse_lease_comparable(block)
            raw["_page_index"] = i
            results.append(raw)
    pdf.close()
    return results

# ============================================================
# FIELD ORDER — matches the example LeaseCompWriteUpExcel.xlsx
# ============================================================

FIELD_ORDER = [
    "Rent Comp",
    "Status",
    "Property Name",
    "Location",
    "City:",
    "State:",
    "Net Rentable Area (SF):",
    "Estimated Year of Construction",
    "Rentable Unit Size (SF):",
    "Tenant",
    "Data Source:",
    "Lease Rate ($/SF/YR):",
    "Adj Lease Rate ($/SF/YR):",
    "Leased Date",
    "Lease Start",
    "Lease Term (months)",
    "Lease End",
    "Months on Market",
    "Months Vacant",
    "Expense Structure",
    "Building Type:",
    "Land Size (SF):",
    "Land Size (Acre):",
    "Additional Comments:",
]


def derive_fields(raw, comp_number=1):
    """Map Collier lease raw fields to the LeaseComp field structure."""
    d = {}

    d["Rent Comp"] = f"No. {comp_number}"
    d["Status"] = "Leased"
    d["Property Name"] = raw.get("name", "")
    d["Location"] = raw.get("address", "")
    d["City:"] = raw.get("city", "")
    d["State:"] = raw.get("state", "Texas")

    nra = raw.get("nra", "")
    d["Net Rentable Area (SF):"] = nra

    d["Estimated Year of Construction"] = raw.get("year_built", "")

    # Rentable Unit Size — use the lease size if available, else NRA
    lease_size = raw.get("lease_size", "")
    d["Rentable Unit Size (SF):"] = lease_size if lease_size else nra

    d["Tenant"] = raw.get("tenant_name", "")

    source = raw.get("source", "")
    company = raw.get("company", "")
    if company and source:
        d["Data Source:"] = f"Colliers International - {source} ({company})"
    elif source:
        d["Data Source:"] = f"Colliers International - {source}"
    else:
        d["Data Source:"] = "Colliers International"

    lease_rate = raw.get("lease_rate", "")
    adj_rate = raw.get("adj_lease_rate", "")
    # Use the contract lease rate (LEASE RATE column), NOT the appraiser-adjusted
    # rate (ADJ LEASE RATE). The adjusted rate is a normalized figure used in the
    # appraisal's comparison analysis and does not reflect actual contract rent.
    d["Lease Rate ($/SF/YR):"] = f"${lease_rate}" if lease_rate else ""
    d["Adj Lease Rate ($/SF/YR):"] = f"${adj_rate}" if adj_rate else ""

    # Dates — the Collier format gives "Current" for start date and term in years
    start_date = raw.get("start_date", "")
    d["Leased Date"] = ""  # Not available in Collier format
    d["Lease Start"] = start_date if start_date != "Current" else ""

    term_years = raw.get("term_years", "")
    if term_years:
        try:
            d["Lease Term (months)"] = str(int(term_years) * 12)
        except ValueError:
            d["Lease Term (months)"] = ""
    else:
        d["Lease Term (months)"] = ""

    d["Lease End"] = ""  # Not directly available; analyst fills from remarks

    d["Months on Market"] = ""  # Not available in Collier format
    d["Months Vacant"] = ""  # Not available

    rate_type = raw.get("rate_type", "")
    d["Expense Structure"] = rate_type

    d["Building Type:"] = f"Convenience Store / Gas Station - {raw.get('construction', '')}" if raw.get("construction") else "Convenience Store / Gas Station"

    site_sf = raw.get("site_sf", "")
    d["Land Size (SF):"] = site_sf
    if site_sf:
        try:
            d["Land Size (Acre):"] = f"{float(site_sf) / 43560:.3f}"
        except ValueError:
            d["Land Size (Acre):"] = ""
    else:
        d["Land Size (Acre):"] = ""

    d["Additional Comments:"] = raw.get("remarks", "")

    return d

# ============================================================
# EXCEL GENERATION (matches example format)
# ============================================================

def _border(style="thin"):
    s = Side(border_style=style, color="000000")
    return Border(left=s, right=s, top=s, bottom=s)

def generate_excel(all_comps, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    ws.column_dimensions["A"].width = 2
    ws.column_dimensions["B"].width = 32
    for col in ["C", "D", "E", "F", "G"]:
        ws.column_dimensions[col].width = 24

    hdr_fill   = PatternFill("solid", fgColor="4472C4")
    label_fill = PatternFill("solid", fgColor="D9E1F2")
    alt_fill   = PatternFill("solid", fgColor="F2F2F2")
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    hdr_font   = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    label_font = Font(name="Calibri", size=10)
    val_font   = Font(name="Calibri", size=10)
    center_al  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_al    = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    # Row 1 — blank (matches example)
    # Row 2+ — data starting with Rent Comp header

    for col_idx, comp in enumerate(all_comps, start=3):
        cell = ws.cell(row=2, column=col_idx)
        cell.value     = comp.get("Rent Comp", f"No. {col_idx - 2}")
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = center_al
        cell.border    = _border()

    # Data rows — starting from row 2 to match example format
    for row_offset, field in enumerate(FIELD_ORDER):
        row_idx  = row_offset + 2  # start at row 2
        use_alt  = (row_offset % 2 == 0)
        row_fill = alt_fill if use_alt else white_fill

        label_cell = ws.cell(row=row_idx, column=2)
        label_cell.value     = field.rstrip(":").rstrip()
        label_cell.font      = label_font
        label_cell.fill      = label_fill
        label_cell.alignment = left_al
        label_cell.border    = _border()

        for ci, comp in enumerate(all_comps):
            val_cell = ws.cell(row=row_idx, column=ci + 3)
            val_cell.font      = val_font
            val_cell.fill      = row_fill
            val_cell.alignment = left_al
            val_cell.border    = _border()

            val = comp.get(field, "")
            # Write numeric where appropriate
            for numeric_field in ("Net Rentable Area (SF):", "Rentable Unit Size (SF):",
                                  "Land Size (SF):", "Lease Term (months)"):
                if field == numeric_field and val:
                    try:
                        val_cell.value = int(float(val))
                        break
                    except (ValueError, TypeError):
                        pass
            else:
                val_cell.value = val

    ws.freeze_panes = "B2"
    try:
        wb.save(output_path)
        print(f"  Excel saved: {output_path}")
    except PermissionError:
        alt = output_path.replace(".xlsx", " (new).xlsx")
        wb.save(alt)
        print(f"  Original file locked (open in Excel?). Saved to: {alt}")

# ============================================================
# PAGE IMAGE RENDERING (JPEG — no orphaned blobs)
# ============================================================

def render_page_as_jpeg(pdf_path, page_index):
    """Render a single PDF page as compressed JPEG bytes."""
    pdf = fitz.open(pdf_path)
    page = pdf[page_index]
    scale = RENDER_DPI / 72.0
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)

    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=JPEG_QUALITY)
    buf.seek(0)
    jpeg_bytes = buf.read()

    pdf.close()
    return jpeg_bytes, pix.width, pix.height

# ============================================================
# WORD DOC GENERATION (blank doc, no template baggage)
# ============================================================

def _add_data_table(doc, comp):
    """Add data table showing only non-blank fields."""
    visible = [(f, comp.get(f, "")) for f in FIELD_ORDER if comp.get(f, "")]
    if not visible:
        return

    table = doc.add_table(rows=len(visible), cols=2)
    table.style = "Table Grid"

    for row_idx, (field, value) in enumerate(visible):
        row = table.rows[row_idx]

        label_cell = row.cells[0]
        label_cell.text = field.rstrip(":").rstrip()
        lp = label_cell.paragraphs[0]
        if lp.runs:
            lp.runs[0].font.bold = True
            lp.runs[0].font.size = Pt(9)
            lp.runs[0].font.name = "Calibri"
        label_cell.width = Inches(2.5)

        val_cell = row.cells[1]
        val_cell.text = str(value)
        vp = val_cell.paragraphs[0]
        if vp.runs:
            vp.runs[0].font.size = Pt(9)
            vp.runs[0].font.name = "Calibri"
        val_cell.width = Inches(3.8)

    return table

def generate_docx(all_comps, pdf_path, page_indices, output_path):
    """Generate Word doc with data tables + PDF page images (JPEG).
    
    Uses Document() blank construction — no template loaded as base —
    to avoid orphaned image blobs that cause Word crashes on copy/paste.
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    # Track which pages we've already rendered (multiple comps per page)
    rendered_pages = set()

    for comp_idx, comp in enumerate(all_comps, start=1):
        # Heading
        heading = doc.add_paragraph(style="Normal")
        run = heading.add_run(f"Lease Comparable No. {comp_idx}")
        run.bold      = True
        run.font.size = Pt(12)
        run.font.name = "Calibri"

        # Data table
        _add_data_table(doc, comp)
        doc.add_paragraph()  # spacer

    # Add PDF page images after all data tables (each page rendered once)
    unique_pages = sorted(set(page_indices))
    for page_idx in unique_pages:
        jpeg_bytes, w, h = render_page_as_jpeg(pdf_path, page_idx)
        target_w = Inches(IMAGE_WIDTH_IN)
        target_h = int(target_w * (h / w))
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(io.BytesIO(jpeg_bytes), width=target_w, height=Emu(target_h))
        doc.add_paragraph()

    try:
        doc.save(output_path)
        print(f"  Word doc saved: {output_path}")
    except PermissionError:
        alt = output_path.replace(".docx", " (new).docx")
        doc.save(alt)
        print(f"  Original file locked (open in Word?). Saved to: {alt}")

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate LeaseComp Excel + Word doc from Collier lease comps PDF"
    )
    parser.add_argument("--name", help="Output name (e.g. 'CStore')")
    parser.add_argument("--pdf",  help="Collier lease comps PDF path")
    args = parser.parse_args()

    if args.name:
        OUTPUT_NAME = args.name
    if args.pdf:
        PDF_PATH = args.pdf

    if not os.path.exists(PDF_PATH):
        print(f"ERROR: PDF not found: {PDF_PATH}")
        raise SystemExit(1)

    print(f"Parsing Collier lease PDF: {PDF_PATH}")
    comps_raw = parse_lease_pdf(PDF_PATH)
    print(f"  Found {len(comps_raw)} lease comparable(s)")

    all_comps = []
    page_indices = []
    for i, raw in enumerate(comps_raw, start=1):
        comp = derive_fields(raw, comp_number=i)
        all_comps.append(comp)
        page_indices.append(raw["_page_index"])
        print(f"  Comp {i}: {raw.get('name','?')} - {raw.get('address','?')}, {raw.get('city','?')}")
        print(f"    NRA: {raw.get('nra','')}  Rate: ${raw.get('adj_lease_rate', raw.get('lease_rate',''))}/SF  "
              f"Type: {raw.get('rate_type','')}  Term: {raw.get('term_years','')} yrs")

    print("\nGenerating Excel...")
    xlsx_path = os.path.join(OUTPUT_DIR, f"LeaseComp_{OUTPUT_NAME}.xlsx")
    generate_excel(all_comps, xlsx_path)

    print("\nGenerating Word doc...")
    docx_path = os.path.join(OUTPUT_DIR, f"LeaseComp_{OUTPUT_NAME}.docx")
    generate_docx(all_comps, PDF_PATH, page_indices, docx_path)

    print(f"\nDone. Output in: {OUTPUT_DIR}")
