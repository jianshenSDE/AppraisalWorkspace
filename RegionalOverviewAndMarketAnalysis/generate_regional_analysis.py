"""
Regional Overview and Market Area Analysis — Document Generator
Generates a DOCX matching the exact formatting of the example template.

Usage:
    cd c:\\AppraisalWorkspace\\RegionalOverviewAndMarketAnalysis
    python generate_regional_analysis.py

Output:
    Output\\CorpusChristiRegional Overview and Market Area Analysis.docx
"""

import shutil
import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# PATHS (relative to this script's directory)
# ---------------------------------------------------------------------------
SCRIPT_DIR    = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "Examples", "RegionalOverviewAndMarketAnalysis", "DeepEastTexasRegional Overview and Market Area Analysis - Completed.docx")
WDA_PDF_PATH  = os.path.join(os.path.dirname(SCRIPT_DIR), "Sources", "WDA", "CorpusChristi.pdf")
OUTPUT_DIR    = os.path.join(SCRIPT_DIR, "Output")
OUTPUT_PATH   = os.path.join(OUTPUT_DIR, "CorpusChristiRegional Overview and Market Area Analysis.docx")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Copy template → output (preserves all custom styles exactly)
# ---------------------------------------------------------------------------
shutil.copy(TEMPLATE_PATH, OUTPUT_PATH)
doc = Document(OUTPUT_PATH)

# Remove every body element except section properties
body = doc.element.body
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------
BULLET   = "\u2022"   # •  used in Normal-style inline bullets
EN_DASH  = "\u2013"   # –  used as separator in List Paragraph items
TODO_MAP = "<TODO: Insert Map>"

def add_heading1(doc, text):
    p = doc.add_paragraph(text, style="Style 1 - heading 1")
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph(text, style="Style2 - heading 2")
    return p

def add_normal(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p

def add_subheading(doc, text, size_pt=14):
    """Bold sub-heading inside a Normal paragraph (14pt Calibri by default)."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    return p

def add_map_label(doc, text=TODO_MAP):
    """Map placeholder — 12pt Bold in Normal paragraph."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    return p

def add_land_use_header(doc, text):
    """Bold 10pt sub-sub-header (Residential, Commercial, etc.) in Normal."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = True
    return p

def embed_pdf_pages(doc, pdf_path, page_width_inches=6.362, dpi=150):
    """Render every page of a PDF as an image and insert into doc at given width."""
    tmp_dir = tempfile.mkdtemp()
    pdf = fitz.open(pdf_path)
    try:
        scale = dpi / 72.0
        mat = fitz.Matrix(scale, scale)
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_path = os.path.join(tmp_dir, f"page_{page_num + 1}.png")
            pix.save(img_path)
            para = doc.add_paragraph(style="Normal")
            run = para.add_run()
            run.add_picture(img_path, width=Inches(page_width_inches))
    finally:
        pdf.close()
        for f in os.listdir(tmp_dir):
            os.remove(os.path.join(tmp_dir, f))
        os.rmdir(tmp_dir)


def add_bullet_normal(doc, text):
    """Inline bullet (•) inside a Normal paragraph."""
    p = doc.add_paragraph(style="Normal")
    p.add_run(f"{BULLET} {text}")
    return p

def add_list_item(doc, text):
    """List Paragraph style item."""
    p = doc.add_paragraph(style="List Paragraph")
    p.add_run(text)
    return p

# ---------------------------------------------------------------------------
# ═══════════════════════════════════════════════════════════════════
# PART 1 — REGIONAL OVERVIEW
# ═══════════════════════════════════════════════════════════════════
# ---------------------------------------------------------------------------
add_heading1(doc, "Regional Overview")

# ── Introduction ────────────────────────────────────────────────────────────
add_subheading(doc, "Introduction")

add_normal(doc,
    "The subject is located in the Corpus Christi Metropolitan Statistical Area (MSA) Workforce "
    "Development Area (WDA). The Corpus Christi MSA includes the following three counties: "
    "Aransas, Nueces, and San Patricio."
)
add_normal(doc,
    "The region is coastal in character and situated along the Texas Gulf Coast in South Texas, "
    "approximately 210 miles southwest of Houston and 140 miles southeast of San Antonio. The "
    "primary population and employment center is Corpus Christi, the largest city in the Coastal "
    "Bend region. The regional economy is anchored by energy production and export, port "
    "operations, petrochemical manufacturing, healthcare services, military installations, "
    "tourism, and government-related employment."
)
add_normal(doc,
    "The Corpus Christi MSA functions as a major regional economic hub and the nation's leading "
    "crude oil export gateway, serving both domestic and international markets through the Port "
    "of Corpus Christi. The region benefits from its strategic Gulf Coast location, robust "
    "transportation infrastructure, and a growing industrial base driven by energy sector "
    "investment."
)

# ── Regional Map ────────────────────────────────────────────────────────────
add_heading2(doc, "Regional Map")
add_map_label(doc, TODO_MAP)

# ── Source citation ──────────────────────────────────────────────────────────
add_normal(doc,
    "The following information is from the Corpus Christi MSA Labor Market Profile, "
    "published by the Texas Workforce Commission (December 2025)"
)

# Embed all WDA PDF pages as images (mirrors the example DOCX structure)
embed_pdf_pages(doc, WDA_PDF_PATH)

# ── Labor Force ─────────────────────────────────────────────────────────────
add_subheading(doc, "Labor Force")

add_normal(doc, "As of December 2025, the Corpus Christi MSA had:")
add_bullet_normal(doc, "Civilian Labor Force: 214,696")
add_bullet_normal(doc, "Employed: 205,364")
add_bullet_normal(doc, "Unemployed: 9,332")
add_bullet_normal(doc, "Unemployment Rate: 4.3%")

add_normal(doc,
    "The unemployment rate of 4.3% is above the Texas average of 3.9% and the national average "
    "of 4.1%, indicating a modestly higher level of labor market slack compared to broader state "
    "and national economies. The Corpus Christi MSA labor force grew by approximately 3,599 "
    "workers year-over-year, reflecting continued regional population and employment expansion."
)
add_normal(doc,
    "The estimated average weekly wage for the Corpus Christi MSA in Q3 2025 was approximately "
    "$1,133, derived from aggregate covered wages of approximately $2.96 billion in the quarter. "
    "While below the Texas average of $1,444, wages reflect the region's mix of high-paying "
    "energy and construction positions offset by a significant share of service, hospitality, "
    "and trade employment."
)

# ── Industry Overview ────────────────────────────────────────────────────────
add_subheading(doc, "Industry Overview")

add_normal(doc,
    "The Corpus Christi MSA maintains a diversified employment base anchored by government "
    "services, energy and construction, trade, and healthcare. As of December 2025, the largest "
    "employment sectors included:"
)
add_list_item(doc, f"Government {EN_DASH} 17.7%")
add_list_item(doc, f"Trade, Transportation and Utilities {EN_DASH} 17.4%")
add_list_item(doc, f"Private Education and Health Services {EN_DASH} 17.1%")
add_list_item(doc, f"Leisure and Hospitality {EN_DASH} 13.0%")
add_list_item(doc, f"Mining, Logging and Construction {EN_DASH} 11.9%")
add_list_item(doc, f"Professional and Business Services {EN_DASH} 9.9%")
add_list_item(doc, f"Financial Activities {EN_DASH} 4.6%")
add_list_item(doc, f"Manufacturing {EN_DASH} 4.5%")
add_list_item(doc, f"Other Services {EN_DASH} 3.2%")
add_list_item(doc, f"Information {EN_DASH} 0.7%")

add_normal(doc,
    "Government, Trade/Transportation, and Private Education/Health together account for "
    "approximately 52.2% of total employment, underscoring the region's reliance on "
    "public-sector services, wholesale and retail trade, port-related logistics activity, "
    "and healthcare."
)
add_normal(doc,
    "Mining, Logging and Construction represents a significant 11.9% of employment, "
    "reflecting the ongoing strength of the energy sector and active construction "
    "activity tied to port expansion, petrochemical facility development, and "
    "major infrastructure investment including the Harbor Bridge Replacement Project."
)
add_normal(doc,
    "Leisure and Hospitality accounts for 13.0% of employment, consistent with Corpus "
    "Christi's role as a major Gulf Coast tourism and hospitality destination, supported "
    "by beaches, bay access, and the Texas State Aquarium."
)

# ── Employment Trends and Projection ────────────────────────────────────────
add_subheading(doc, "Employment Trends and Projection")

add_normal(doc,
    "Between 2022 and 2032, the fastest-growing industries in the Corpus Christi MSA are "
    "projected to include:"
)
add_list_item(doc, f"Support Activities for Oil and Gas Extraction {EN_DASH} 22.5%")
add_list_item(doc, f"Individual and Family Services {EN_DASH} 21.3%")
add_list_item(doc, f"Offices of Physicians {EN_DASH} 19.8%")
add_list_item(doc, f"Child Care Services {EN_DASH} 18.4%")
add_list_item(doc, f"Restaurants and Other Eating Places {EN_DASH} 17.5%")
add_list_item(doc, f"Nursing Care Facilities (Skilled Nursing Facilities) {EN_DASH} 16.3%")
add_list_item(doc, f"Leisure and Hospitality {EN_DASH} 15.8%")
add_list_item(doc, f"Specialty Trade Contractors {EN_DASH} 14.2%")
add_list_item(doc, f"Professional and Business Services {EN_DASH} 12.1%")
add_list_item(doc, f"Building Material and Supplies Dealers {EN_DASH} 11.4%")

add_normal(doc,
    "Projected growth reflects the dual drivers of the Corpus Christi economy: continued "
    "energy sector investment and associated construction and support services, alongside "
    "healthcare and social services growth driven by an aging regional population and "
    "expanding medical infrastructure."
)
add_normal(doc,
    "Growth in restaurants, leisure and hospitality, and building material suppliers "
    "indicates continued incremental population growth and ongoing residential and "
    "commercial construction activity across the region, supported by strong in-migration "
    "trends and Gulf Coast tourism demand."
)

# ── Employment by Firm Size and Ownership ───────────────────────────────────
add_subheading(doc, "Employment by Firm Size and Ownership")

add_normal(doc,
    "Employment within the Corpus Christi MSA is concentrated in larger firms, with a "
    "notable presence of major industrial and governmental employers:"
)
add_list_item(doc, "26.1% of jobs are in firms with 1,000+ employees")
add_list_item(doc, "15.6% of jobs are in firms with 100\u2013249 employees")
add_list_item(doc, "13.4% are in firms with 20\u201349 employees")
add_list_item(doc, "11.1% are in firms with 50\u201399 employees")
add_list_item(doc, "Smaller firms (1\u20139 employees) collectively account for approximately 8.4%")

add_normal(doc,
    "The high concentration of large-firm employment (26.1% in firms with 1,000+ "
    "employees) reflects the presence of major energy companies, port operations, "
    "hospital systems, military installations, and government agencies that are "
    "characteristic of this coastal industrial economy."
)
add_normal(doc,
    "Government employment accounts for approximately 17.7% of total nonfarm employment, "
    "with private sector employment comprising approximately 82.3% of total jobs. The "
    "government sector includes federal installations such as Naval Air Station Corpus "
    "Christi, state agencies, and local government and public school districts."
)

# ── Unemployment Trends ──────────────────────────────────────────────────────
add_subheading(doc, "Unemployment Trends")

add_normal(doc,
    "Historical data indicates that unemployment in the Corpus Christi MSA peaked during "
    "the COVID-19 pandemic in April 2020 at approximately 14.5%, before declining through "
    "the remainder of 2020 and steadily improving through 2021 and 2022."
)
add_normal(doc,
    "From 2023 through 2025, unemployment has generally stabilized within the 3.9% to "
    "5.0% range, with periodic seasonal fluctuations consistent with the region's "
    "tourism and construction employment patterns."
)
add_normal(doc,
    "As of December 2025, unemployment stands at 4.3%, reflecting a stable labor market "
    "that remains modestly above the Texas average of 3.9% but demonstrates consistent "
    "improvement from post-pandemic levels."
)

# ── Conclusion (Regional Overview) ──────────────────────────────────────────
add_subheading(doc, "Conclusion")

add_normal(doc,
    "The Corpus Christi MSA represents a dynamic coastal economy anchored by energy "
    "production and export, port operations, healthcare, government, and tourism, "
    "serving a labor force of approximately 215,000 workers across Aransas, Nueces, "
    "and San Patricio counties. With the Port of Corpus Christi ranking as the nation's "
    "leading crude oil export gateway and over $65 billion in capital investment "
    "flowing into the port and surrounding industrial ecosystem, the region demonstrates "
    "strong long-term economic fundamentals."
)
add_normal(doc,
    "Key strengths include energy and petrochemical operations, a major deepwater port, "
    "healthcare services, military installations, regional retail trade, and a growing "
    "tourism and hospitality sector. Projected growth in energy support services, "
    "healthcare, and construction reflects both demographic demand and continued "
    "industrial expansion."
)
add_normal(doc,
    "While unemployment remains modestly above the Texas average and wages trail the "
    "statewide benchmark, the Corpus Christi MSA demonstrates continued economic "
    "resilience supported by substantial capital investment, essential energy "
    "infrastructure, and a strategic Gulf Coast location with national and "
    "international market access."
)

# ---------------------------------------------------------------------------
# ═══════════════════════════════════════════════════════════════════
# PART 2 — MARKET AREA ANALYSIS
# ═══════════════════════════════════════════════════════════════════
# ---------------------------------------------------------------------------
add_heading1(doc, "Market Area Analysis")

add_normal(doc,
    "A market area is the geographic area in which the subject property competes for "
    "the attentions of market participants; the term broadly defines an area containing "
    "diverse land uses. Market areas are defined by a combination of factors including "
    "physical features, and socioeconomic characteristics of the residents or tenants, "
    "the condition of the improvements and land use trends. Market area analysis focuses "
    "on the identification of boundaries and the social, economic, governmental and "
    "environmental influences that affect the value of real property within those "
    "boundaries."
)
add_normal(doc,
    "The purpose of a market area analysis is to provide a bridge between the study of "
    "general influences on all property values and the analysis of a particular subject. "
    "Market area boundaries are identified by determining the area in which the four "
    "forces that affect value (social, economic, governmental and environmental) operate "
    "in the same way they affect the subject property. Interaction of the various "
    "components influencing these four forces often results in the dissimilarities "
    "regarding the length of time between the stages of a market area\u2019s life cycle."
)

# ── Market area Map ──────────────────────────────────────────────────────────
add_map_label(doc, "Market area Map")
add_map_label(doc, TODO_MAP)

# ── General Description ──────────────────────────────────────────────────────
add_subheading(doc, "General Description")

add_normal(doc,
    "For this analysis, the market area is defined as Aransas County, Texas. Aransas "
    "County is located along the Texas Gulf Coast approximately 30 miles northeast of "
    "Corpus Christi and approximately 200 miles southwest of Houston. The county seat "
    "is Rockport, which serves as the primary governmental, commercial, and employment "
    "center of the county. Other communities within the county include Fulton, Lamar, "
    "Holiday Beach, and Rockport Beach, while the city of Aransas Pass straddles the "
    "county line with San Patricio County."
)
add_normal(doc,
    "Aransas County encompasses a total area of approximately 528 square miles, of which "
    "approximately 276 square miles (52%) are covered by water, including Aransas Bay, "
    "Copano Bay, and portions of the Intracoastal Waterway. The county is situated on "
    "the Live Oak Peninsula, a narrow coastal landmass bordered by Copano Bay to the "
    "west and Aransas Bay to the east, with direct access to the Gulf of Mexico via "
    "Aransas Pass and Port Aransas."
)
add_normal(doc,
    "The county is characterized by a coastal landscape of bay shoreline, wetlands, "
    "salt marshes, live oak mottes, and open water that supports a thriving commercial "
    "fishing, recreational boating, and eco-tourism economy. The Aransas National "
    "Wildlife Refuge, located partially within the county, is internationally recognized "
    "as the wintering habitat of the endangered whooping crane and draws significant "
    "tourism activity annually."
)
add_normal(doc,
    "Aransas County functions as a coastal resort and retirement community with a growing "
    "residential base, serving both permanent residents and seasonal visitors from across "
    "Texas and the broader region. The county\u2019s proximity to the Corpus Christi MSA "
    "provides access to regional employment, healthcare, and commercial services while "
    "maintaining its distinctive small-town coastal character."
)

# ── Access And Major Roadways ─────────────────────────────────────────────────
add_subheading(doc, "Access And Major Roadways")

add_normal(doc,
    "Aransas County is served by several state highways that provide connectivity to "
    "surrounding employment centers, regional markets, and the Corpus Christi MSA:"
)
add_bullet_normal(doc,
    "State Highway 35 \u2013 The primary north-south arterial running through Rockport and "
    "the county, connecting Aransas County to Portland and Corpus Christi to the south "
    "and to Refugio, Victoria, and the Houston metro region to the north. SH 35 serves "
    "as the main commercial and transportation corridor through the county."
)
add_bullet_normal(doc,
    "State Highway 188 \u2013 Provides access from SH 35 eastward through the Live Oak "
    "Peninsula toward the Aransas National Wildlife Refuge and the Lamar area, serving "
    "rural residential, agricultural, and wildlife-related land uses in the eastern "
    "portion of the county."
)
add_bullet_normal(doc,
    "State Highway 361 \u2013 Connects the county to Port Aransas via Aransas Pass and "
    "Ingleside in San Patricio County to the west, and serves as the access route to "
    "the ferry crossing of Corpus Christi Channel to Port Aransas and Mustang Island. "
    "SH 361 supports tourism traffic and provides an alternate route to Corpus Christi "
    "via the barrier islands."
)
add_bullet_normal(doc,
    "Farm-to-Market Roads \u2013 Numerous FM roads and county-maintained roadways provide "
    "access to rural residential properties, waterfront developments, agricultural "
    "tracts, and coastal recreational areas throughout the county."
)
add_bullet_normal(doc,
    "Aransas County Airport \u2013 Located in Fulton, north of Rockport, the airport "
    "provides general aviation services for private and charter aircraft, supporting "
    "regional business and recreational access."
)
add_normal(doc,
    "While Aransas County does not have direct interstate highway access, the regional "
    "highway network provides adequate transportation connectivity for residents, "
    "commerce, and coastal tourism. These roadways facilitate commuting access to the "
    "Corpus Christi MSA employment base and support the movement of commercial fishing "
    "products, petroleum-related traffic, and recreational tourism throughout the "
    "county and surrounding region."
)

# ── Land Use And Supportive Development ──────────────────────────────────────
add_subheading(doc, "Land Use And Supportive Development")

add_normal(doc,
    "Aransas County exhibits a coastal land use pattern with development concentrated "
    "primarily in Rockport, Fulton, and areas along the Aransas and Copano Bay "
    "shorelines, with significant portions of the county devoted to water, wetlands, "
    "and protected natural areas."
)

add_land_use_header(doc, "Residential")
add_normal(doc,
    "Residential development is concentrated in Rockport and Fulton along the bay "
    "waterfront and adjacent inland areas, with rural residential properties scattered "
    "throughout the peninsula. The county maintains approximately 15,500 housing units, "
    "with approximately 34% vacant \u2014 a high vacancy rate reflecting the significant "
    "number of seasonal and vacation homes that characterize this coastal resort community. "
    "Among occupied units, approximately 77.5% are owner-occupied. Residential development "
    "ranges from modest single-family homes and manufactured housing to upscale waterfront "
    "estates, condominium complexes, and second-home subdivisions. The county\u2019s "
    "median age of 52.5 years and 28.7% of residents aged 65 or older reflects its "
    "appeal as a retirement and second-home destination."
)

add_land_use_header(doc, "Commercial")
add_normal(doc,
    "Commercial development is concentrated along SH 35 and Business SH 35-L through "
    "Rockport, where retail stores, restaurants, marinas, boat repair facilities, "
    "medical offices, financial institutions, hotels, and tourism-oriented businesses "
    "serve both the local population and the substantial visitor trade. Rockport\u2019s "
    "downtown historic district along Fulton Beach Road and the Rockport Beach area "
    "contain additional commercial and hospitality establishments catering to the "
    "tourism economy. Commercial activity reflects a dual market: year-round services "
    "for permanent residents and seasonal retail, dining, and lodging for visitors."
)

add_land_use_header(doc, "Industrial")
add_normal(doc,
    "Industrial land uses within Aransas County are limited but include commercial "
    "fishing operations, seafood processing facilities, boat building and repair "
    "yards, marine supply operations, and petroleum pipeline facilities associated "
    "with the regional energy infrastructure. The Port of Rockport supports commercial "
    "fishing and recreational boating activity. Light industrial and marine-related "
    "businesses are interspersed along the waterfront and bay access areas. The county "
    "does not contain significant heavy manufacturing or large-scale industrial "
    "operations."
)

add_land_use_header(doc, "Agricultural and Open Land")
add_normal(doc,
    "Agriculture and open land represent meaningful land uses in the inland and "
    "peninsular areas of Aransas County. Ranching, hay production, and limited row "
    "crop agriculture occur on upland tracts. Coastal wetlands, tidal flats, and "
    "marshlands cover extensive areas particularly in the eastern portion of the county "
    "adjacent to the Aransas National Wildlife Refuge. Large tracts of open land are "
    "also utilized for recreational hunting, wildlife observation, and eco-tourism "
    "activities that leverage the county\u2019s exceptional biodiversity and natural "
    "coastal resources."
)

add_land_use_header(doc, "Recreational and Institutional")
add_normal(doc,
    "Aransas County contains significant recreational and institutional land uses "
    "including the Rockport-Fulton Independent School District facilities, county "
    "government buildings, churches, Rockport Veterans Memorial Park, and the Rockport "
    "Center for the Arts. Recreational resources include Rockport Beach Park, Goose "
    "Island State Park, the Aransas National Wildlife Refuge, numerous public boat "
    "ramps, and extensive bay and Gulf of Mexico fishing access. These resources support "
    "the county\u2019s significant nature-based and outdoor recreation tourism economy, "
    "drawing hundreds of thousands of visitors annually for fishing, birding, kayaking, "
    "and coastal sightseeing."
)

# ── Life Stage and Trends ─────────────────────────────────────────────────────
add_subheading(doc, "Life Stage and Trends")

add_normal(doc,
    "The market area is considered to be in a stable to gradually expanding stage "
    "with steady long-term growth characteristics driven primarily by in-migration, "
    "retirement demand, and Gulf Coast tourism. Aransas County\u2019s population "
    "has shown moderate growth consistent with coastal South Texas communities "
    "experiencing increased demand from retirees, second-home buyers, and Gulf "
    "Coast lifestyle-seekers."
)
add_normal(doc,
    "Economic stability and growth are supported by:"
)
add_bullet_normal(doc, "Gulf Coast tourism, fishing, and outdoor recreation")
add_bullet_normal(doc, "Retirement and second-home residential demand")
add_bullet_normal(doc, "Commercial fishing and marine industries")
add_bullet_normal(doc, "Proximity to Corpus Christi MSA employment base")
add_bullet_normal(doc, "Environmental and nature tourism (whooping cranes, birding, kayaking)")
add_normal(doc,
    "The county experienced significant disruption from Hurricane Harvey in August 2017, "
    "which inflicted extensive damage on residential and commercial properties throughout "
    "Aransas County. Post-Harvey recovery and rebuilding activity generated substantial "
    "construction employment and resulted in significant residential and commercial "
    "reinvestment, with many properties rebuilt to higher standards. The recovery "
    "period generated meaningful new development activity and infrastructure improvement "
    "across the county."
)
add_normal(doc,
    "Recent development patterns reflect continued residential demand, particularly "
    "for waterfront and bay-view properties, single-family homes, and vacation rental "
    "units. Commercial development has expanded to support the growing tourism and "
    "hospitality sector, with new hotel, restaurant, and marina developments adding "
    "to the county\u2019s visitor amenities."
)
add_normal(doc,
    "Existing improvements throughout the county reflect a range of property ages and "
    "conditions, from post-Harvey new construction to older pre-storm structures. "
    "Post-Harvey rebuilding has introduced modern construction standards and elevated "
    "building requirements across much of the county."
)
add_normal(doc,
    "Future development trends are expected to include:"
)
add_bullet_normal(doc, "Continued coastal and waterfront residential development")
add_bullet_normal(doc, "Growth in short-term vacation rental and hospitality properties")
add_bullet_normal(doc, "Ongoing commercial expansion along SH 35 and the Rockport waterfront")
add_bullet_normal(doc, "Eco-tourism and nature-based recreational facility development")
add_normal(doc,
    "Overall growth is expected to remain steady and consistent with the county\u2019s "
    "coastal resort and retirement market dynamics, supported by continued in-migration "
    "and the sustained appeal of Gulf Coast living."
)

# ── Conclusion (Market Area) ──────────────────────────────────────────────────
add_subheading(doc, "Conclusion")

add_normal(doc,
    "Aransas County provides a stable and gradually expanding coastal market environment "
    "characterized by bay and Gulf waterfront development, coastal resort and retirement "
    "residential uses, nature-based tourism, commercial fishing, and marine-related "
    "industries centered in Rockport."
)
add_normal(doc, "Key factors supporting the local market include:")
add_list_item(doc, "Regional highway connectivity via SH 35 and access to Corpus Christi MSA")
add_list_item(doc, "Gulf Coast location supporting tourism, fishing, and outdoor recreation")
add_list_item(doc, "Retirement and second-home demand driving residential growth")
add_list_item(doc, "Aransas National Wildlife Refuge and natural resources supporting eco-tourism")
add_list_item(doc, "Post-Hurricane Harvey recovery and reinvestment in improved structures")
add_list_item(doc, "Proximity to Port of Corpus Christi and regional energy economy")
add_normal(doc,
    "Although development intensity is lower than in metropolitan regions, Aransas County "
    "demonstrates stable long-term land use patterns and steady demand for coastal "
    "residential, recreational, and commercial properties. Future growth is expected to "
    "remain gradual and market-driven, with development focused around existing "
    "communities, waterfront areas, and major transportation corridors."
)
add_normal(doc,
    "Overall, Aransas County represents a stable coastal market area with consistent "
    "land use patterns and long-term viability supported by Gulf Coast tourism, "
    "retirement in-migration, commercial fishing, and regional connectivity to the "
    "broader Corpus Christi MSA economic base."
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUTPUT_PATH)
print(f"Document saved to:\n  {OUTPUT_PATH}")
