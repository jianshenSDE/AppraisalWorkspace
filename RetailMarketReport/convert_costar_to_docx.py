"""
Convert a CoStar Retail Market Report PDF into a formatted Word document.

Usage:
    python convert_costar_to_docx.py <input_pdf> [--output <output_docx>] [--template <template_docx>]

The script:
  1. Parses the CoStar PDF to identify sections (Overview, Leasing, Rent, etc.)
  2. Classifies each page as narrative text or chart/table using font analysis
  3. Extracts narrative text from two-column pages and reformats into paragraphs
  4. Renders chart/table pages as high-resolution images
  5. Produces a Word document matching the template formatting exactly

Requirements:
    pip install python-docx PyMuPDF Pillow lxml
"""

import argparse
import io
import os
import re
import sys
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "example",
                             "Austin MSA Retail Market Report.docx")

# Canonical section names in document order
SECTION_ORDER = [
    "Overview",
    "Leasing",
    "Rent",
    "Construction",
    "Under Construction Properties",
    "Sales",
    "Sales Past 12 Months",
    "Economy",
    "Submarkets",
    "Supply & Demand Trends",
    "Rent & Vacancy",
    "Sale Trends",
]

# How sections are named in the Word doc headings (some differ from PDF)
SECTION_DISPLAY_NAMES = {
    "Overview": "Overview",
    "Leasing": "Leasing",
    "Rent": "Rent",
    "Construction": "Construction",
    "Under Construction Properties": "Under Construction Properties",
    "Sales": "Sales",
    "Sales Past 12 Months": "Sales Past 12 Months",
    "Economy": "Economy",
    "Submarkets": "Submarkets",
    "Supply & Demand Trends": "Supply and Demand Trends",
    "Rent & Vacancy": "Rent & Vacancy",
    "Sale Trends": "Sale Trends",
}

# DPI for rendering PDF pages as images
RENDER_DPI = 200

# Image width in the Word doc (6.50 inches = within 1-inch margins on 8.5" page)
IMAGE_WIDTH_EMU = int(6.50 * 914400)

# Footer crop: CoStar footer (date, copyright, page number) starts at ~y=742
# on a 792pt-high page.  We clip at y=738 to remove it cleanly.
FOOTER_CROP_Y = 738


# ---------------------------------------------------------------------------
# PDF Page Classification
# ---------------------------------------------------------------------------

def classify_page_by_font(pdf_doc, page_idx):
    """Classify a page as 'narrative' or 'image' using font analysis.

    Narrative pages have substantial body text in Arial ~10pt (non-bold).
    Table/chart pages have mostly small/bold text or very little body text.

    Returns 'narrative' if >50% of body characters are narrative-style AND
    there are at least 200 such characters. Otherwise returns 'image'.
    """
    page = pdf_doc[page_idx]
    page_dict = page.get_text("dict")
    page_height = page.rect.height

    narrative_chars = 0
    other_chars = 0

    for block in page_dict["blocks"]:
        if block["type"] != 0:
            continue
        y0 = block["bbox"][1]
        # Skip header bar (top ~60pt) and footer (bottom ~60pt)
        if y0 < 60 or y0 > page_height - 60:
            continue

        for line in block["lines"]:
            for span in line["spans"]:
                font = span["font"]
                size = span["size"]
                text = span["text"].strip()
                if not text:
                    continue

                if size >= 20:
                    # Large metric numbers
                    other_chars += len(text)
                elif size >= 14:
                    # Section sub-headers
                    other_chars += len(text)
                elif 9.5 <= size <= 11 and "Bold" not in font:
                    # Body narrative text (Arial ~10pt, not bold)
                    narrative_chars += len(text)
                else:
                    other_chars += len(text)

    total = narrative_chars + other_chars
    if total == 0:
        return "image"
    narrative_pct = narrative_chars / total * 100
    return "narrative" if narrative_pct > 50 and narrative_chars > 200 else "image"


# ---------------------------------------------------------------------------
# PDF Section Parsing
# ---------------------------------------------------------------------------

def identify_sections(pdf_doc):
    """Parse the PDF and group pages by section.

    Returns list of dicts:
        [{"name": "Overview", "pages": [2, 3]}, ...]
    Skips cover/TOC pages (typically pages 0 and 1).
    """
    sections = []
    current_section = None
    current_pages = []

    for i in range(len(pdf_doc)):
        page = pdf_doc[i]
        text = page.get_text("text").strip()
        first_line = text.split('\n')[0].strip() if text else ""

        # Check if this page starts a new (or same) section
        matched = None
        for sec_name in SECTION_ORDER:
            if first_line == sec_name:
                matched = sec_name
                break

        if matched and matched != current_section:
            if current_section is not None:
                sections.append({"name": current_section, "pages": current_pages})
            current_section = matched
            current_pages = [i]
        elif current_section is not None:
            current_pages.append(i)
        # else: cover/TOC pages — skip

    if current_section is not None:
        sections.append({"name": current_section, "pages": current_pages})

    return sections


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

def extract_narrative_paragraphs(pdf_doc, page_idx):
    """Extract narrative paragraphs from a two-column PDF page.

    Uses font analysis to select only body text (Arial ~10pt, non-bold),
    splits into left/right columns, reads left then right, and joins
    lines within each text block into paragraphs.
    """
    page = pdf_doc[page_idx]
    page_width = page.rect.width
    page_height = page.rect.height
    mid_x = page_width / 2

    blocks = page.get_text("dict")["blocks"]

    # Collect body text blocks with their bounding boxes
    body_blocks = []  # (x0, y0, x1, y1, paragraph_text)

    for block in blocks:
        if block["type"] != 0:
            continue
        x0, y0, x1, y1 = block["bbox"]
        # Skip header/footer
        if y0 < 60 or y0 > page_height - 60:
            continue

        # Check if this block is narrative text
        narrative_spans = []
        total_chars = 0
        narrative_chars = 0
        for line in block["lines"]:
            for span in line["spans"]:
                text = span["text"].strip()
                if not text:
                    continue
                size = span["size"]
                font = span["font"]
                total_chars += len(text)
                if 9.5 <= size <= 11 and "Bold" not in font:
                    narrative_chars += len(text)
                    narrative_spans.append(span)

        # Only include blocks that are predominantly narrative text
        if total_chars > 0 and narrative_chars / total_chars > 0.7 and narrative_chars > 30:
            # Reconstruct paragraph text from narrative spans
            lines_text = []
            current_line_y = None
            current_line_parts = []
            for line in block["lines"]:
                line_y = line["bbox"][1]
                line_text_parts = []
                for span in line["spans"]:
                    text = span["text"]
                    size = span["size"]
                    font = span["font"]
                    if 9.5 <= size <= 11 and "Bold" not in font:
                        line_text_parts.append(text)
                if line_text_parts:
                    lines_text.append(''.join(line_text_parts).strip())

            paragraph_text = ' '.join(lines_text)
            if paragraph_text.strip():
                body_blocks.append((x0, y0, x1, y1, paragraph_text.strip()))

    if not body_blocks:
        return []

    # Split into left and right column
    left = [(x0, y0, x1, y1, t) for x0, y0, x1, y1, t in body_blocks if x0 < mid_x - 10]
    right = [(x0, y0, x1, y1, t) for x0, y0, x1, y1, t in body_blocks if x0 >= mid_x - 10]

    # Sort by vertical position
    left.sort(key=lambda b: b[1])
    right.sort(key=lambda b: b[1])

    # Check if last left block and first right block form a continuation
    # (i.e., a paragraph that spans across columns)
    paragraphs = []

    # Process left column
    for _, _, _, _, text in left:
        paragraphs.append(text)

    # Check for cross-column paragraph continuation:
    # If the last left paragraph ends without a period and the first right
    # paragraph starts with a lowercase letter, they're likely one paragraph.
    if left and right:
        last_left = paragraphs[-1] if paragraphs else ""
        first_right = right[0][4]
        if (last_left and not last_left.endswith('.')
                and first_right and first_right[0].islower()):
            # Merge them
            paragraphs[-1] = last_left + ' ' + first_right
            right = right[1:]

    # Process right column
    for _, _, _, _, text in right:
        paragraphs.append(text)

    return paragraphs


def render_metrics_bar(pdf_doc, page_idx, dpi=RENDER_DPI):
    """Render the metrics summary bar from the top of a section's first page.

    The CoStar pages have a metrics bar (big bold numbers + labels)
    roughly between y=55 and y=140 (in PDF points).
    Returns PNG bytes, or None if no metrics bar detected.
    """
    page = pdf_doc[page_idx]
    page_dict = page.get_text("dict")

    # Check for large metric numbers in the header area
    has_metrics = False
    for block in page_dict["blocks"]:
        if block["type"] != 0:
            continue
        y0 = block["bbox"][1]
        if 55 < y0 < 145:
            for line in block["lines"]:
                for span in line["spans"]:
                    if span["size"] >= 20 and any(c.isdigit() for c in span["text"]):
                        has_metrics = True
                        break
                if has_metrics:
                    break
        if has_metrics:
            break

    if not has_metrics:
        return None

    zoom = dpi / 72.0
    clip = fitz.Rect(0, 55, page.rect.width, 140)
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
    return pix.tobytes("png")


def render_page_as_image(pdf_doc, page_idx, dpi=RENDER_DPI):
    """Render a PDF page as PNG bytes, cropping out the CoStar footer."""
    page = pdf_doc[page_idx]
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    clip = fitz.Rect(0, 0, page.rect.width, FOOTER_CROP_Y)
    pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
    return pix.tobytes("png")


# ---------------------------------------------------------------------------
# PDF Metadata Extraction
# ---------------------------------------------------------------------------

def get_msa_name(pdf_doc):
    """Extract MSA name from page 1 (e.g., 'Austin Retail' → 'Austin')."""
    page = pdf_doc[1] if len(pdf_doc) > 1 else pdf_doc[0]
    text = page.get_text("text").strip()
    for line in text.split('\n'):
        line = line.strip()
        if line.endswith('Retail'):
            return line.replace('Retail', '').strip()
    return "Unknown MSA"


def get_costar_date(pdf_doc):
    """Extract report date from PDF footer (e.g., '3/1/2026')."""
    for i in range(min(5, len(pdf_doc))):
        for line in pdf_doc[i].get_text("text").split('\n'):
            match = re.match(r'^\s*(\d{1,2}/\d{1,2}/\d{4})\s*$', line)
            if match:
                return match.group(1)
    return "Unknown Date"


def format_date_long(date_str):
    """Convert '3/1/2026' → 'March 1, 2026'."""
    from datetime import datetime
    try:
        dt = datetime.strptime(date_str, "%m/%d/%Y")
        # Avoid zero-padded day
        return dt.strftime("%B ") + str(dt.day) + dt.strftime(", %Y")
    except ValueError:
        return date_str


# ---------------------------------------------------------------------------
# Word Document Generation
# ---------------------------------------------------------------------------

def add_image_paragraph(doc, img_data, center=True):
    """Add a full-width image in a centered paragraph, preserving aspect ratio."""
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Calculate height to preserve aspect ratio
    img = Image.open(io.BytesIO(img_data))
    w, h = img.size
    aspect = h / w
    target_h = int(IMAGE_WIDTH_EMU * aspect)

    run = para.add_run()
    run.add_picture(io.BytesIO(img_data), width=Emu(IMAGE_WIDTH_EMU), height=Emu(target_h))


def create_report(pdf_path, output_path, template_path=None):
    """Create a Word document from a CoStar Retail Market Report PDF.

    Args:
        pdf_path:     Path to the CoStar PDF.
        output_path:  Path for the output .docx.
        template_path: Path to the template .docx (provides styles).
    """
    if template_path is None:
        template_path = TEMPLATE_PATH

    print(f"Opening PDF:  {pdf_path}")
    pdf_doc = fitz.open(pdf_path)

    print(f"Template:     {template_path}")
    doc = Document(template_path)

    # Clear all content from the template, keeping section properties (margins, etc.)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)

    # Extract metadata
    msa_name = get_msa_name(pdf_doc)
    costar_date = get_costar_date(pdf_doc)
    formatted_date = format_date_long(costar_date)
    sections = identify_sections(pdf_doc)

    print(f"MSA:          {msa_name}")
    print(f"Date:         {costar_date} -> {formatted_date}")
    print(f"Sections:     {len(sections)}")

    # --- Title ---
    doc.add_paragraph(f"{msa_name} MSA Retail Market Report",
                      style='Style 1 - heading 1')

    # --- Source line (italic) ---
    source_para = doc.add_paragraph(style='No Spacing')
    source_run = source_para.add_run(f"Source: Costar as of {formatted_date}")
    source_run.italic = True

    # --- Process each section ---
    for sec in sections:
        sec_name = sec["name"]
        page_indices = sec["pages"]
        display = SECTION_DISPLAY_NAMES.get(sec_name, sec_name)

        # Classify every page in this section
        narrative_pages = []
        image_pages = []
        for pg in page_indices:
            pg_type = classify_page_by_font(pdf_doc, pg)
            if pg_type == "narrative":
                narrative_pages.append(pg)
            else:
                image_pages.append(pg)

        has_narrative = len(narrative_pages) > 0
        print(f"  {sec_name:40s}  pages={len(page_indices)}  "
              f"narrative={len(narrative_pages)}  image={len(image_pages)}")

        # --- Add section heading ---
        # In the example doc, most sections use Style2-heading2.
        # "Rent & Vacancy" and "Sale Trends" appear as plain Normal text in the example.
        if sec_name in ("Rent & Vacancy", "Sale Trends"):
            p = doc.add_paragraph(display, style='Normal')
        else:
            p = doc.add_paragraph(display)
            p.style = doc.styles['Style2 - heading 2']

        # --- Metrics bar image (for narrative sections' first page) ---
        if has_narrative:
            first_narrative = narrative_pages[0]
            metrics_img = render_metrics_bar(pdf_doc, first_narrative)
            if metrics_img:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Calculate proper height
                pil_img = Image.open(io.BytesIO(metrics_img))
                w, h = pil_img.size
                target_h = int(IMAGE_WIDTH_EMU * (h / w))
                run = img_para.add_run()
                run.add_picture(io.BytesIO(metrics_img),
                                width=Emu(IMAGE_WIDTH_EMU),
                                height=Emu(target_h))

        # --- Narrative text paragraphs ---
        if has_narrative:
            all_paragraphs = []
            for pg in narrative_pages:
                paras = extract_narrative_paragraphs(pdf_doc, pg)
                all_paragraphs.extend(paras)
            for para_text in all_paragraphs:
                doc.add_paragraph(para_text, style='Normal')

        # --- Image pages (charts, tables, etc.) ---
        # These are rendered in page order (they come after narrative pages
        # in the section).  We render them in the order they appear in the PDF.
        all_image_pages = sorted(image_pages)
        for pg in all_image_pages:
            img_data = render_page_as_image(pdf_doc, pg)
            add_image_paragraph(doc, img_data)

    # --- Save ---
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    pdf_doc.close()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Convert a CoStar Retail Market Report PDF to a Word document."
    )
    parser.add_argument("input_pdf", help="Path to the CoStar PDF")
    parser.add_argument("--output", "-o",
                        help="Output .docx path (default: auto-generated in output/)")
    parser.add_argument("--template", "-t",
                        help="Template .docx path (default: example template)",
                        default=TEMPLATE_PATH)

    args = parser.parse_args()
    input_pdf = os.path.abspath(args.input_pdf)

    if not os.path.exists(input_pdf):
        print(f"Error: Input PDF not found: {input_pdf}")
        sys.exit(1)

    if args.output:
        output_path = os.path.abspath(args.output)
    else:
        pdf_name = Path(input_pdf).stem
        match = re.match(r'^(.+?)\s*-\s*TX\s+USA', pdf_name)
        msa = match.group(1).strip() if match else pdf_name
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{msa} MSA Retail Market Report.docx")

    create_report(input_pdf, output_path, args.template)


if __name__ == "__main__":
    main()
